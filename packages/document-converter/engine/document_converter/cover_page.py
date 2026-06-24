"""Cover page generator for DOCX documents."""
import re
from pathlib import Path
from typing import List, Optional
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


def get_min_heading_level(markdown_content: str) -> int:
    """Get the minimum heading level used in the markdown content.

    Args:
        markdown_content: The markdown content to parse.

    Returns:
        Minimum heading level (1-6), or 0 if no headings found.
    """
    heading_pattern = re.compile(r'^(#{1,6})\s+.+$', re.MULTILINE)
    matches = heading_pattern.findall(markdown_content)
    if not matches:
        return 0
    return min(len(h) for h in matches)


def normalize_heading_levels(markdown_content: str) -> str:
    """Normalize heading levels so that the minimum level becomes H1.

    If markdown only contains H2+ headings (no H1), this promotes all headings
    so the minimum level becomes H1. For example:
    - If min is H2: H2→H1, H3→H2, H4→H3, etc.
    - If min is H3: H3→H1, H4→H2, H5→H3, etc.

    Args:
        markdown_content: The markdown content to modify.

    Returns:
        Modified markdown with normalized heading levels.
    """
    min_level = get_min_heading_level(markdown_content)

    # If no headings or already has H1, no normalization needed
    if min_level <= 1:
        return markdown_content

    # Calculate how many levels to promote
    levels_to_promote = min_level - 1

    lines = markdown_content.split('\n')
    result = []

    for line in lines:
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            hashes = match.group(1)
            text = match.group(2)
            current_level = len(hashes)
            new_level = current_level - levels_to_promote
            # Ensure we don't go below H1
            new_level = max(1, new_level)
            line = '#' * new_level + ' ' + text
        result.append(line)

    return '\n'.join(result)


def count_h1_headings(markdown_content: str) -> int:
    """Count the number of H1 headings in markdown content.

    Args:
        markdown_content: The markdown content to parse.

    Returns:
        Number of H1 headings found.
    """
    h1_pattern = re.compile(r'^#\s+.+$', re.MULTILINE)
    return len(h1_pattern.findall(markdown_content))


def should_create_cover_page(markdown_content: str) -> bool:
    """Determine if a cover page should be created.

    A cover page is only created when there is exactly one H1 heading.

    Args:
        markdown_content: The markdown content to parse.

    Returns:
        True if cover page should be created, False otherwise.
    """
    return count_h1_headings(markdown_content) == 1


def demote_headings(markdown_content: str) -> str:
    """Demote heading levels appropriately.

    The original H1 line is removed (it goes to cover page).
    Content under H1 (before first H2) is kept.
    - H3/H4/etc before first H2 → H1 (promoted to top level)
    - H2 → H1
    - H3 after first H2 → H2
    - etc.

    Args:
        markdown_content: The markdown content to modify.

    Returns:
        Modified markdown with adjusted headings.
    """
    lines = markdown_content.split('\n')
    result = []
    h1_found = False
    first_h2_found = False

    for line in lines:
        # Skip the H1 line only (it goes to cover page)
        if re.match(r'^#\s+[^#]', line) and not h1_found:
            h1_found = True
            continue

        # Check if this is the first H2
        if re.match(r'^##\s+[^#]', line) and not first_h2_found:
            first_h2_found = True

        # Handle headings based on position
        if not first_h2_found:
            # Before first H2: promote H3+ to H1
            if re.match(r'^###+\s+', line):
                line = re.sub(r'^###+\s+', '# ', line)
        else:
            # After first H2: demote normally (## -> #, ### -> ##, etc.)
            if re.match(r'^######\s+', line):
                line = re.sub(r'^######\s+', '##### ', line)
            elif re.match(r'^#####\s+', line):
                line = re.sub(r'^#####\s+', '#### ', line)
            elif re.match(r'^####\s+', line):
                line = re.sub(r'^####\s+', '### ', line)
            elif re.match(r'^###\s+', line):
                line = re.sub(r'^###\s+', '## ', line)
            elif re.match(r'^##\s+', line):
                line = re.sub(r'^##\s+', '# ', line)

        result.append(line)

    return '\n'.join(result)


def extract_title_and_description(markdown_content: str) -> tuple[str, str, bool]:
    """Extract the H1 title and description from markdown content.

    The description is any content between the H1 and the first H2.
    Only extracts if there is exactly one H1 heading.

    Args:
        markdown_content: The markdown content to parse.

    Returns:
        Tuple of (title, description, should_create_cover).
        Empty strings and False if cover page should not be created.
    """
    # Only create cover page if exactly one H1
    if not should_create_cover_page(markdown_content):
        return '', '', False
    # Find H1 title (# Title)
    h1_match = re.search(r'^#\s+(.+?)$', markdown_content, re.MULTILINE)
    if not h1_match:
        return '', '', False

    title = h1_match.group(1).strip()

    # Find content between H1 and first H2
    h1_end = h1_match.end()
    h2_match = re.search(r'^##\s+', markdown_content[h1_end:], re.MULTILINE)

    if h2_match:
        description_text = markdown_content[h1_end:h1_end + h2_match.start()]
    else:
        description_text = markdown_content[h1_end:]

    # Clean up description - remove markdown formatting, keep text
    description = description_text.strip()
    # Remove ** bold markers but keep the text
    description = re.sub(r'\*\*(.+?)\*\*', r'\1', description)
    # Remove heading markers (### Header -> Header)
    description = re.sub(r'^#{1,6}\s+', '', description, flags=re.MULTILINE)
    # Remove bullet point markers (* item -> item, - item -> item)
    description = re.sub(r'^\s*[\*\-]\s+', '', description, flags=re.MULTILINE)
    # Remove horizontal rules
    description = re.sub(r'^---+$', '', description, flags=re.MULTILINE)
    # Clean up multiple blank lines
    description = re.sub(r'\n{3,}', '\n\n', description)
    description = description.strip()

    return title, description, True


def create_page_break_paragraph():
    """Create a paragraph with a page break.

    Returns:
        A paragraph element containing a page break.
    """
    paragraph = OxmlElement('w:p')
    run = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run.append(br)
    paragraph.append(run)
    return paragraph


def add_cover_page(
    docx_path: Path,
    title: str,
    description: str = ""
) -> None:
    """Add a cover page to the beginning of a DOCX document.

    The cover page contains:
    - Centered H1 title
    - Description text below the title
    - Page break after the cover page

    Args:
        docx_path: Path to the DOCX file.
        title: The main title for the cover page.
        description: Optional description text below the title.
    """
    if not title:
        return

    doc = Document(docx_path)
    body = doc.element.body

    # Find the first element
    first_element = body[0] if len(body) > 0 else None

    # Create cover page elements (in reverse order for insertion)
    elements_to_insert = []

    # Page break after cover page
    page_break = create_page_break_paragraph()
    elements_to_insert.append(page_break)

    # Description paragraphs (if any)
    if description:
        for line in reversed(description.split('\n')):
            line = line.strip()
            if line:
                desc_para = OxmlElement('w:p')

                # Paragraph properties - center align with spacing
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)

                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:after'), '120')  # 6pt after
                pPr.append(spacing)

                desc_para.append(pPr)

                # Run with text
                run = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '24')  # 12pt
                rPr.append(sz)
                run.append(rPr)

                text_elem = OxmlElement('w:t')
                text_elem.set(qn('xml:space'), 'preserve')
                text_elem.text = line
                run.append(text_elem)

                desc_para.append(run)
                elements_to_insert.append(desc_para)

    # Spacer before description
    spacer = OxmlElement('w:p')
    spacer_pPr = OxmlElement('w:pPr')
    spacer_spacing = OxmlElement('w:spacing')
    spacer_spacing.set(qn('w:after'), '480')  # 24pt after
    spacer_pPr.append(spacer_spacing)
    spacer.append(spacer_pPr)
    elements_to_insert.append(spacer)

    # Title paragraph - centered H1 style
    title_para = OxmlElement('w:p')

    # Paragraph properties
    pPr = OxmlElement('w:pPr')

    # Center alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)

    # Spacing before (to push title down on page)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '4800')  # ~2.5 inches from top
    spacing.set(qn('w:after'), '240')    # 12pt after
    pPr.append(spacing)

    title_para.append(pPr)

    # Run with title text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Bold
    bold = OxmlElement('w:b')
    rPr.append(bold)

    # Font size - 36pt for cover title
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '72')  # 36pt (half-points)
    rPr.append(sz)

    # Color - dark blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2F5496')
    rPr.append(color)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = title
    run.append(text_elem)

    title_para.append(run)
    elements_to_insert.append(title_para)

    # Insert all elements at the beginning
    for elem in elements_to_insert:
        if first_element is not None:
            first_element.addprevious(elem)
        else:
            body.append(elem)
        first_element = elem

    # Remove the original H1 from the document body (it's now on cover page)
    # Find and remove the first Heading 1 paragraph
    for para in doc.paragraphs:
        if para.style and para.style.name == 'Heading 1':
            # Check if this is our title
            if para.text.strip() == title:
                p_element = para._element
                p_element.getparent().remove(p_element)
                break

    doc.save(docx_path)


def copy_template_cover_page(
    output_path: Path,
    template_path: Path
) -> bool:
    """Copy the cover page from a template to the output document.

    This copies all elements from the template up to and including the first
    page break, preserving textboxes, drawings, and other complex elements.

    If the template has no page break (single-page cover template), all template
    elements are copied as the cover page and a page break is added after.

    Args:
        output_path: Path to the output DOCX file.
        template_path: Path to the template DOCX file.

    Returns:
        True if cover page was copied, False if template has no content.
    """
    template_doc = Document(template_path)
    output_doc = Document(output_path)

    template_body = template_doc.element.body
    output_body = output_doc.element.body

    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Find all elements up to and including the first page break
    cover_elements = []
    found_page_break = False

    for elem in template_body:
        # Deep copy the element
        elem_copy = deepcopy(elem)
        cover_elements.append(elem_copy)

        # Check if this element contains a page break
        for br in elem.iter(f'{w_ns}br'):
            if br.get(qn('w:type')) == 'page':
                found_page_break = True
                break

        if found_page_break:
            break

    if not cover_elements:
        return False

    # Get the first element in output to insert before
    first_output = output_body[0] if len(output_body) > 0 else None

    # Insert cover page elements at the beginning of output document
    for elem in reversed(cover_elements):
        if first_output is not None:
            first_output.addprevious(elem)
        else:
            output_body.append(elem)
        first_output = elem

    # If template had no page break, we need to add one after the cover page
    # This ensures the content starts on a new page after the cover
    if not found_page_break:
        # Create a page break paragraph
        page_break_para = create_page_break_paragraph()
        # Insert after the last cover element (which is now at the beginning)
        # cover_elements[0] after reversal is at the end, cover_elements[-1] is at start
        # The first_output now points to the first cover element
        # We need to insert the page break after the last cover element
        last_cover_in_output = cover_elements[-1]  # This was inserted first, so it's last
        # Actually, after insertion, the order is: cover_elements[0], cover_elements[1], ... , cover_elements[-1]
        # first_output now points to cover_elements[0]
        # We need to find the last cover element and insert page break after it

        # Find where to insert the page break (after all cover elements, before content)
        # The safest way is to find the first non-cover element
        content_start = None
        for i, child in enumerate(output_body):
            # Skip cover elements
            if i >= len(cover_elements):
                content_start = child
                break

        if content_start is not None:
            content_start.addprevious(page_break_para)
        else:
            output_body.append(page_break_para)

    output_doc.save(output_path)
    return True  # Always return True if we copied something
