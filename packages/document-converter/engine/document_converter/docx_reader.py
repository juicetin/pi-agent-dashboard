"""DOCX to Markdown extraction module."""
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Any
from dataclasses import dataclass, field
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
import re


@dataclass
class ExtractedElement:
    """Represents an extracted document element."""
    type: str  # 'heading', 'paragraph', 'list_item', 'table', 'image'
    content: str
    level: int = 0  # For headings (1-6) or list nesting
    style_name: str = ""
    docx_path: str = ""  # XPath-like reference to original element
    line_start: int = 0
    line_end: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractionResult:
    """Result of DOCX extraction."""
    markdown: str
    elements: List[ExtractedElement]
    preserved_elements: List[Dict[str, Any]]
    images: List[Dict[str, Any]]


def get_paragraph_style_level(paragraph: Paragraph) -> Tuple[str, int]:
    """Determine the type and level of a paragraph based on its style.

    Args:
        paragraph: The paragraph to analyze.

    Returns:
        Tuple of (element_type, level) where element_type is 'heading',
        'list_item', or 'paragraph'.
    """
    style_name = paragraph.style.name if paragraph.style else ""

    # Check for heading styles
    if style_name.startswith('Heading'):
        try:
            level = int(style_name.split()[-1])
            return ('heading', min(level, 6))
        except (ValueError, IndexError):
            pass

    # Check for Title style (treat as H1)
    if style_name == 'Title':
        return ('heading', 1)

    # Check for Subtitle style (treat as H2)
    if style_name == 'Subtitle':
        return ('heading', 2)

    # Check for list styles
    if 'List' in style_name or 'Bullet' in style_name:
        return ('list_item', 1)

    return ('paragraph', 0)


def get_list_info(paragraph: Paragraph) -> Tuple[bool, bool, int]:
    """Check if paragraph is a list item and get its properties.

    Args:
        paragraph: The paragraph to check.

    Returns:
        Tuple of (is_list, is_numbered, indent_level).
    """
    # Check for numbering in the paragraph's XML
    pPr = paragraph._element.pPr
    if pPr is None:
        return (False, False, 0)

    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return (False, False, 0)

    # Get indent level
    ilvl = numPr.find(qn('w:ilvl'))
    indent_level = int(ilvl.get(qn('w:val'))) if ilvl is not None else 0

    # Check if numbered (vs bullet)
    numId = numPr.find(qn('w:numId'))
    is_numbered = False
    if numId is not None:
        # NumId > 0 indicates a list; we'd need to check the numbering definitions
        # for now, assume non-zero numId with certain patterns are numbered
        num_val = numId.get(qn('w:val'))
        # This is a simplification - proper detection requires checking abstractNum
        is_numbered = num_val is not None and int(num_val) > 0

    return (True, is_numbered, indent_level)


def parse_paragraph(paragraph: Paragraph, index: int) -> ExtractedElement:
    """Parse a single paragraph into an ExtractedElement.

    Args:
        paragraph: The paragraph to parse.
        index: The paragraph index in the document.

    Returns:
        ExtractedElement with parsed content.
    """
    text = paragraph.text.strip()
    style_name = paragraph.style.name if paragraph.style else ""
    element_type, level = get_paragraph_style_level(paragraph)

    # Check for list formatting
    is_list, is_numbered, indent_level = get_list_info(paragraph)
    if is_list:
        element_type = 'list_item'
        level = indent_level

    return ExtractedElement(
        type=element_type,
        content=text,
        level=level,
        style_name=style_name,
        docx_path=f"/w:body/w:p[{index + 1}]",
        metadata={
            'is_numbered': is_numbered if is_list else None,
            'indent_level': indent_level if is_list else None,
        }
    )


def parse_table(table: Table, index: int) -> ExtractedElement:
    """Parse a table into Markdown format.

    Args:
        table: The table to parse.
        index: The table index in the document.

    Returns:
        ExtractedElement with Markdown table content.
    """
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ExtractedElement(
            type='table',
            content='',
            docx_path=f"/w:body/w:tbl[{index + 1}]"
        )

    # Build Markdown table
    md_lines = []

    # Header row
    header = rows[0]
    md_lines.append('| ' + ' | '.join(header) + ' |')

    # Separator
    md_lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')

    # Data rows
    for row in rows[1:]:
        # Pad row if necessary
        while len(row) < len(header):
            row.append('')
        md_lines.append('| ' + ' | '.join(row[:len(header)]) + ' |')

    return ExtractedElement(
        type='table',
        content='\n'.join(md_lines),
        docx_path=f"/w:body/w:tbl[{index + 1}]",
        metadata={'rows': len(rows), 'cols': len(header) if header else 0}
    )


def extract_table_style(table: Table) -> Dict[str, Any]:
    """Extract style information from a table.

    Args:
        table: The table to extract styles from.

    Returns:
        Dictionary with table style properties.
    """
    style: Dict[str, Any] = {}

    # Try to get table properties
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))

    if tbl_pr is not None:
        # Check for table style reference
        tbl_style = tbl_pr.find(qn('w:tblStyle'))
        if tbl_style is not None:
            style['style_name'] = tbl_style.get(qn('w:val'))

    # Extract styles from first row (header)
    if table.rows:
        header_row = table.rows[0]
        for cell in header_row.cells:
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                # Check for shading (background color)
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill != 'auto':
                        style['header_bg_color'] = fill

                # Check for borders
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    # Get border properties from any border
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.find(qn(f'w:{border_name}'))
                        if border is not None:
                            border_color = border.get(qn('w:color'))
                            if border_color:
                                style['border_color'] = border_color
                            border_size = border.get(qn('w:sz'))
                            if border_size:
                                style['border_size'] = int(border_size)
                            break

            # Check for text color in header
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.color.rgb:
                        # Convert RGBColor to hex string
                        rgb = run.font.color.rgb
                        style['header_text_color'] = f"{rgb}"
                    break
                break
            break

    # Check for alternating row colors (striping)
    if len(table.rows) > 2:
        row1_color = None
        row2_color = None

        # Check second row
        if len(table.rows) > 1:
            for cell in table.rows[1].cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        row1_color = shd.get(qn('w:fill'))
                break

        # Check third row
        if len(table.rows) > 2:
            for cell in table.rows[2].cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        row2_color = shd.get(qn('w:fill'))
                break

        # If colors alternate, it's striped
        if row1_color and row2_color and row1_color != row2_color:
            style['striped'] = True
            style['odd_row_color'] = row1_color
            style['even_row_color'] = row2_color
        elif row1_color:
            style['striped'] = False

    return style


def extract_table_styles(doc: DocxDocument) -> Dict[str, Any]:
    """Extract default table style from document.

    Args:
        doc: The document to extract styles from.

    Returns:
        Dictionary with default table style (from most common style).
    """
    if not doc.tables:
        return {}

    # Collect styles from all tables
    all_styles = []
    for table in doc.tables:
        style = extract_table_style(table)
        if style:
            all_styles.append(style)

    if not all_styles:
        return {}

    # Return most common style properties
    # For simplicity, use the first table's style as default
    return all_styles[0] if all_styles else {}


def extract_per_table_styles(doc: DocxDocument) -> List[Dict[str, Any]]:
    """Extract per-table style overrides.

    Args:
        doc: The document to analyze.

    Returns:
        List of style dictionaries for tables that differ from default.
    """
    if not doc.tables:
        return []

    # Get default style from first table
    default_style = extract_table_style(doc.tables[0]) if doc.tables else {}

    per_table_styles = []
    for i, table in enumerate(doc.tables):
        table_style = extract_table_style(table)

        # Check if this table differs from default
        diff_style = {}
        for key, value in table_style.items():
            if default_style.get(key) != value:
                diff_style[key] = value

        if diff_style:
            diff_style['table_index'] = i
            per_table_styles.append(diff_style)

    return per_table_styles


def calculate_column_widths(table: Table) -> Dict[str, float]:
    """Calculate column width multipliers for a table.

    Args:
        table: The table to analyze.

    Returns:
        Dictionary mapping column header text to width multipliers.
    """
    if not table.rows:
        return {}

    header_row = table.rows[0]
    widths: Dict[str, float] = {}
    total_width = 0
    cell_widths = []

    # Get widths from each header cell
    for cell in header_row.cells:
        cell_text = cell.text.strip()
        tcPr = cell._tc.find(qn('w:tcPr'))

        width = 1.0  # Default width
        if tcPr is not None:
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is not None:
                w_val = tcW.get(qn('w:w'))
                if w_val and w_val.isdigit():
                    width = int(w_val)

        cell_widths.append((cell_text, width))
        total_width += width

    # Normalize to multipliers (relative to average)
    if total_width > 0 and cell_widths:
        avg_width = total_width / len(cell_widths)
        for cell_text, width in cell_widths:
            if cell_text:  # Only include cells with text
                widths[cell_text] = round(width / avg_width, 2)

    return widths


def find_images(doc: DocxDocument) -> List[Dict[str, Any]]:
    """Find all images in the document.

    Args:
        doc: The document to search.

    Returns:
        List of image information dictionaries.
    """
    images = []

    # Find all drawing elements
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            # Check for inline shapes (images)
            drawing_elements = run._element.findall('.//' + qn('w:drawing'))
            for j, drawing in enumerate(drawing_elements):
                # Look for blip elements which contain image references
                blips = drawing.findall('.//' + qn('a:blip'))
                for blip in blips:
                    embed = blip.get(qn('r:embed'))
                    if embed:
                        images.append({
                            'type': 'image',
                            'embed_id': embed,
                            'paragraph_index': i,
                            'docx_path': f"/w:body/w:p[{i + 1}]/w:r/w:drawing[{j + 1}]"
                        })

    return images


def find_preserved_elements(doc: DocxDocument) -> List[Dict[str, Any]]:
    """Find elements that should be preserved (charts, shapes, etc).

    Args:
        doc: The document to search.

    Returns:
        List of preserved element information.
    """
    preserved = []

    # Find charts
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            drawing_elements = run._element.findall('.//' + qn('w:drawing'))
            for j, drawing in enumerate(drawing_elements):
                # Check for chart references
                charts = drawing.findall('.//' + qn('c:chart'))
                for chart in charts:
                    chart_id = chart.get(qn('r:id'))
                    if chart_id:
                        preserved.append({
                            'type': 'chart',
                            'ref_id': chart_id,
                            'paragraph_index': i,
                            'docx_path': f"/w:body/w:p[{i + 1}]/w:r/w:drawing[{j + 1}]/chart"
                        })

    return preserved


def element_to_markdown(element: ExtractedElement, prev_element: Optional[ExtractedElement] = None) -> str:
    """Convert an ExtractedElement to Markdown.

    Args:
        element: The element to convert.
        prev_element: The previous element (for context).

    Returns:
        Markdown string.
    """
    if element.type == 'heading':
        prefix = '#' * element.level
        return f"{prefix} {element.content}"

    elif element.type == 'list_item':
        indent = '  ' * element.level
        is_numbered = element.metadata.get('is_numbered', False)
        if is_numbered:
            return f"{indent}1. {element.content}"
        else:
            return f"{indent}- {element.content}"

    elif element.type == 'table':
        return element.content

    elif element.type == 'image':
        # Image placeholder - actual path set during extraction
        return element.content

    else:  # paragraph
        return element.content


def extract_content(doc_path: Path) -> ExtractionResult:
    """Extract content from a DOCX file.

    Args:
        doc_path: Path to the DOCX file.

    Returns:
        ExtractionResult with markdown content and element mappings.
    """
    doc = Document(doc_path)
    elements: List[ExtractedElement] = []

    # Track indices for different element types
    para_index = 0
    table_index = 0

    # Iterate through document body
    for child in doc.element.body:
        tag = child.tag.split('}')[-1]  # Get tag name without namespace

        if tag == 'p':  # Paragraph
            para = doc.paragraphs[para_index]
            if para.text.strip():  # Only include non-empty paragraphs
                element = parse_paragraph(para, para_index)
                elements.append(element)
            para_index += 1

        elif tag == 'tbl':  # Table
            table = doc.tables[table_index]
            element = parse_table(table, table_index)
            if element.content:
                elements.append(element)
            table_index += 1

    # Find images and preserved elements
    images = find_images(doc)
    preserved = find_preserved_elements(doc)

    # Convert to Markdown with line tracking
    md_lines = []
    current_line = 1
    prev_element = None
    in_list = False

    for element in elements:
        # Add blank line before headings (except at start)
        if element.type == 'heading' and md_lines:
            md_lines.append('')
            current_line += 1

        # Add blank line before tables
        if element.type == 'table' and md_lines:
            md_lines.append('')
            current_line += 1

        # Add blank line after lists when switching to non-list
        if prev_element and prev_element.type == 'list_item' and element.type != 'list_item':
            md_lines.append('')
            current_line += 1
            in_list = False

        # Add blank line before first list item
        if element.type == 'list_item' and not in_list and md_lines:
            md_lines.append('')
            current_line += 1
            in_list = True

        # Track line numbers
        element.line_start = current_line

        # Convert to markdown
        md_content = element_to_markdown(element, prev_element)
        content_lines = md_content.split('\n')
        md_lines.extend(content_lines)

        element.line_end = current_line + len(content_lines) - 1
        current_line = element.line_end + 1

        # Add blank line after headings
        if element.type == 'heading':
            md_lines.append('')
            current_line += 1

        # Add blank line after tables
        if element.type == 'table':
            md_lines.append('')
            current_line += 1

        prev_element = element

    markdown = '\n'.join(md_lines)

    return ExtractionResult(
        markdown=markdown,
        elements=elements,
        preserved_elements=preserved,
        images=images
    )


def convert_docx_to_md(
    input_path: Path,
    output_path: Path,
    include_meta: bool = True,
    extract_images: bool = True
) -> ExtractionResult:
    """Convert a DOCX file to Markdown.

    This is the main entry point for DOCX to Markdown conversion.

    Args:
        input_path: Path to the DOCX file.
        output_path: Path for the output Markdown file.
        include_meta: Whether to generate metadata XML file.
        extract_images: Whether to extract images to files.

    Returns:
        ExtractionResult with extraction details.
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    # Validate input
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if not input_path.suffix.lower() == '.docx':
        raise ValueError(f"Input file must be a .docx file: {input_path}")

    # Extract content
    result = extract_content(input_path)

    # Handle image extraction
    if extract_images and result.images:
        from .image_extractor import extract_images as do_extract_images
        result = do_extract_images(input_path, output_path, result)

    # Write markdown output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(result.markdown, encoding='utf-8')

    # Generate metadata if requested
    if include_meta:
        from .meta_mapper import save_metadata
        meta_path = output_path.parent / f"{output_path.stem}_meta.xml"
        save_metadata(result, input_path, output_path, meta_path)

    return result
