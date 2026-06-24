"""Round-trip editing module for DOCX-Markdown synchronization."""
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from difflib import SequenceMatcher
import shutil

from docx import Document
from docx.oxml.ns import qn

from .meta_mapper import load_metadata, find_metadata_path
from .docx_reader import ExtractionResult


@dataclass
class ChangeInfo:
    """Information about a detected change."""
    mapping_id: str
    change_type: str  # 'modified', 'deleted', 'added'
    old_content: str
    new_content: str
    line_start: int
    line_end: int
    docx_path: str


@dataclass
class MergeResult:
    """Result of a merge operation."""
    success: bool
    output_path: Path
    changes_applied: List[ChangeInfo]
    warnings: List[str]
    preserved_count: int


def detect_changes(
    original_md: str,
    edited_md: str,
    metadata: Dict[str, Any]
) -> List[ChangeInfo]:
    """Detect changes between original and edited Markdown.

    Args:
        original_md: Original Markdown content.
        edited_md: Edited Markdown content.
        metadata: Loaded metadata with mappings.

    Returns:
        List of ChangeInfo objects describing the changes.
    """
    changes = []

    original_lines = original_md.split('\n')
    edited_lines = edited_md.split('\n')

    # Use SequenceMatcher to find differences
    matcher = SequenceMatcher(None, original_lines, edited_lines)

    # Track which mappings are affected
    for mapping in metadata.get('mappings', []):
        start = mapping.get('md_start', 1) - 1  # Convert to 0-indexed
        end = mapping.get('md_end', 1) - 1

        if start < 0:
            continue

        # Get original content for this mapping
        if end < len(original_lines):
            original_content = '\n'.join(original_lines[start:end + 1])
        else:
            original_content = '\n'.join(original_lines[start:])

        # Check if this range was modified
        # Find corresponding range in edited content
        modified = False
        new_content = original_content

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                continue

            # Check if this operation affects our range
            if (i1 <= start <= i2) or (i1 <= end <= i2) or (start <= i1 and end >= i2):
                modified = True

                # Try to find new content
                if tag == 'replace':
                    # Content was modified
                    if j1 < len(edited_lines) and j2 <= len(edited_lines):
                        new_content = '\n'.join(edited_lines[j1:j2])
                elif tag == 'delete':
                    # Content was deleted
                    new_content = ''
                elif tag == 'insert':
                    # New content was added
                    if j1 < len(edited_lines) and j2 <= len(edited_lines):
                        new_content = '\n'.join(edited_lines[j1:j2])

                break

        if modified:
            change_type = 'deleted' if not new_content else 'modified'
            changes.append(ChangeInfo(
                mapping_id=mapping.get('id', ''),
                change_type=change_type,
                old_content=original_content,
                new_content=new_content,
                line_start=start + 1,
                line_end=end + 1,
                docx_path=mapping.get('docx_path', '')
            ))

    return changes


def update_paragraph_content(doc: Document, docx_path: str, new_content: str) -> bool:
    """Update a paragraph's content in the document.

    Args:
        doc: The document to modify.
        docx_path: XPath-like path to the paragraph.
        new_content: New content for the paragraph.

    Returns:
        True if update was successful, False otherwise.
    """
    # Parse the docx_path to find the paragraph
    # Expected format: /w:body/w:p[N]
    import re
    match = re.match(r'/w:body/w:p\[(\d+)\]', docx_path)
    if not match:
        return False

    para_index = int(match.group(1)) - 1  # Convert to 0-indexed

    if para_index < 0 or para_index >= len(doc.paragraphs):
        return False

    para = doc.paragraphs[para_index]

    # Clear existing content
    for run in para.runs:
        run.text = ''

    # Add new content
    if para.runs:
        para.runs[0].text = new_content
    else:
        para.add_run(new_content)

    return True


def restore_preserved_elements(
    doc: Document,
    original_doc: Document,
    preserved: List[Dict[str, Any]]
) -> int:
    """Restore preserved elements from original document.

    Args:
        doc: The document being updated.
        original_doc: The original document.
        preserved: List of preserved element information.

    Returns:
        Number of elements restored.
    """
    restored = 0

    # For now, this is a placeholder
    # Full implementation would copy XML elements from original_doc to doc
    # based on the preserved element paths

    # Charts and complex objects require copying the relationship
    # and the XML element, which is complex with python-docx

    return restored


def merge_document(
    original_docx: Path,
    edited_md: Path,
    output_path: Path,
    meta_path: Optional[Path] = None
) -> MergeResult:
    """Merge edited Markdown back into the original DOCX.

    Args:
        original_docx: Path to the original DOCX file.
        edited_md: Path to the edited Markdown file.
        output_path: Path for the output DOCX file.
        meta_path: Path to the metadata file (auto-detected if None).

    Returns:
        MergeResult with merge details.
    """
    original_docx = Path(original_docx)
    edited_md = Path(edited_md)
    output_path = Path(output_path)

    warnings = []

    # Validate inputs
    if not original_docx.exists():
        raise FileNotFoundError(f"Original DOCX not found: {original_docx}")
    if not edited_md.exists():
        raise FileNotFoundError(f"Edited Markdown not found: {edited_md}")

    # Find metadata file
    if meta_path is None:
        meta_path = find_metadata_path(edited_md)
        if meta_path is None:
            # Try to find it based on markdown filename
            meta_path = edited_md.parent / f"{edited_md.stem}_meta.xml"

    if meta_path is None or not meta_path.exists():
        raise FileNotFoundError(
            f"Metadata file not found. Expected: {edited_md.stem}_meta.xml"
        )

    # Load metadata
    metadata = load_metadata(meta_path)

    # Read original markdown (if available) and edited markdown
    original_md_path = edited_md  # Assume same path for original
    if 'extracted' in metadata:
        original_md_path = edited_md.parent / metadata['extracted']

    # Read the edited markdown
    edited_content = edited_md.read_text(encoding='utf-8')

    # For comparison, we need the original markdown
    # This should ideally be stored or reconstructed
    # For now, we'll do a simplified approach
    original_content = edited_content  # Placeholder

    # Detect changes
    changes = detect_changes(original_content, edited_content, metadata)

    if not changes:
        # No changes detected, just copy the original
        shutil.copy2(original_docx, output_path)
        return MergeResult(
            success=True,
            output_path=output_path,
            changes_applied=[],
            warnings=["No changes detected, copied original document."],
            preserved_count=0
        )

    # Copy original DOCX to output
    shutil.copy2(original_docx, output_path)

    # Open the copy and apply changes
    doc = Document(output_path)
    original_doc = Document(original_docx)

    changes_applied = []
    for change in changes:
        if change.change_type == 'deleted':
            warnings.append(f"Skipping deleted content at {change.docx_path}")
            continue

        success = update_paragraph_content(doc, change.docx_path, change.new_content)
        if success:
            changes_applied.append(change)
        else:
            warnings.append(f"Failed to update {change.docx_path}")

    # Restore preserved elements
    preserved = metadata.get('preserved', [])
    preserved_count = restore_preserved_elements(doc, original_doc, preserved)

    # Save the merged document
    doc.save(output_path)

    return MergeResult(
        success=True,
        output_path=output_path,
        changes_applied=changes_applied,
        warnings=warnings,
        preserved_count=preserved_count
    )


def merge_md_to_docx(
    original_docx: Path,
    edited_md: Path,
    output_path: Path,
    meta_path: Optional[Path] = None
) -> MergeResult:
    """Public API for merging edited Markdown back into DOCX.

    This is the main entry point for round-trip editing.

    Args:
        original_docx: Path to the original DOCX file.
        edited_md: Path to the edited Markdown file.
        output_path: Path for the output DOCX file.
        meta_path: Path to the metadata file (auto-detected if None).

    Returns:
        MergeResult with merge details.
    """
    return merge_document(
        original_docx=Path(original_docx),
        edited_md=Path(edited_md),
        output_path=Path(output_path),
        meta_path=Path(meta_path) if meta_path else None
    )
