"""Style extraction module for DOCX documents."""
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field
from xml.etree import ElementTree as ET
from xml.dom import minidom
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


@dataclass
class StyleInfo:
    """Information about a document style."""
    name: str
    font_name: Optional[str] = None
    font_size: Optional[int] = None  # In half-points (Word internal unit)
    bold: bool = False
    italic: bool = False
    underline: bool = False
    color: Optional[str] = None  # Hex color without #
    background_color: Optional[str] = None
    spacing_before: Optional[int] = None  # In twips
    spacing_after: Optional[int] = None
    line_spacing: Optional[float] = None
    alignment: Optional[str] = None  # left, center, right, justify
    indent_left: Optional[int] = None
    indent_right: Optional[int] = None
    indent_first_line: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class TableStyleInfo:
    """Information about a table style."""
    name: str
    header_fill_color: Optional[str] = None
    header_text_color: Optional[str] = None
    header_bold: bool = True
    row_band_color: Optional[str] = None  # Alternating row color
    border_color: Optional[str] = None
    border_width: Optional[int] = None
    cell_padding: Optional[int] = None


def extract_paragraph_style(style) -> StyleInfo:
    """Extract style information from a paragraph style.

    Args:
        style: The docx style object.

    Returns:
        StyleInfo containing the extracted information.
    """
    info = StyleInfo(name=style.name)

    # Get font information
    font = style.font
    if font:
        if font.name:
            info.font_name = font.name
        if font.size:
            info.font_size = font.size.pt * 2  # Convert to half-points
        info.bold = font.bold or False
        info.italic = font.italic or False
        info.underline = font.underline or False
        if font.color and font.color.rgb:
            info.color = str(font.color.rgb)

    # Get paragraph formatting
    pf = style.paragraph_format
    if pf:
        if pf.space_before:
            info.spacing_before = pf.space_before.twips
        if pf.space_after:
            info.spacing_after = pf.space_after.twips
        if pf.line_spacing:
            info.line_spacing = pf.line_spacing
        if pf.alignment is not None:
            alignment_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
            info.alignment = alignment_map.get(pf.alignment, 'left')
        if pf.left_indent:
            info.indent_left = pf.left_indent.twips
        if pf.right_indent:
            info.indent_right = pf.right_indent.twips
        if pf.first_line_indent:
            info.indent_first_line = pf.first_line_indent.twips

    return info


def extract_styles(doc_path: Path) -> Dict[str, StyleInfo]:
    """Extract all styles from a DOCX document.

    Args:
        doc_path: Path to the DOCX file.

    Returns:
        Dictionary mapping style names to StyleInfo objects.
    """
    doc = Document(doc_path)
    styles = {}

    # Extract paragraph styles
    for style in doc.styles:
        if style.type == 1:  # Paragraph style
            try:
                info = extract_paragraph_style(style)
                styles[style.name] = info
            except Exception:
                # Skip styles that can't be parsed
                pass

    return styles


def extract_table_styles(doc_path: Path) -> List[TableStyleInfo]:
    """Extract table style information from a document.

    Args:
        doc_path: Path to the DOCX file.

    Returns:
        List of TableStyleInfo objects.
    """
    doc = Document(doc_path)
    table_styles = []

    for i, table in enumerate(doc.tables):
        style_info = TableStyleInfo(name=f"Table_{i + 1}")

        # Try to get header row styling
        if table.rows:
            first_row = table.rows[0]
            for cell in first_row.cells:
                # Check for shading (fill color)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill != 'auto':
                        style_info.header_fill_color = fill

                # Check text formatting in header
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.bold:
                            style_info.header_bold = True
                        if run.font.color and run.font.color.rgb:
                            style_info.header_text_color = str(run.font.color.rgb)
                break  # Only check first cell

        table_styles.append(style_info)

    return table_styles


def style_to_xml(style: StyleInfo) -> ET.Element:
    """Convert a StyleInfo to an XML element.

    Args:
        style: The StyleInfo to convert.

    Returns:
        XML Element representing the style.
    """
    elem = ET.Element('style')
    elem.set('name', style.name)

    # Font element
    if style.font_name or style.font_size:
        font = ET.SubElement(elem, 'font')
        if style.font_name:
            font.set('name', style.font_name)
        if style.font_size:
            font.set('size', str(style.font_size))
        if style.bold:
            font.set('bold', 'true')
        if style.italic:
            font.set('italic', 'true')
        if style.underline:
            font.set('underline', 'true')

    # Color
    if style.color:
        color = ET.SubElement(elem, 'color')
        color.set('value', style.color)

    # Background color
    if style.background_color:
        fill = ET.SubElement(elem, 'fill')
        fill.set('color', style.background_color)

    # Spacing
    if style.spacing_before is not None or style.spacing_after is not None:
        spacing = ET.SubElement(elem, 'spacing')
        if style.spacing_before is not None:
            spacing.set('before', str(style.spacing_before))
        if style.spacing_after is not None:
            spacing.set('after', str(style.spacing_after))
        if style.line_spacing is not None:
            spacing.set('line', str(style.line_spacing))

    # Alignment
    if style.alignment:
        align = ET.SubElement(elem, 'alignment')
        align.set('value', style.alignment)

    # Indentation
    if style.indent_left or style.indent_right or style.indent_first_line:
        indent = ET.SubElement(elem, 'indent')
        if style.indent_left:
            indent.set('left', str(style.indent_left))
        if style.indent_right:
            indent.set('right', str(style.indent_right))
        if style.indent_first_line:
            indent.set('first-line', str(style.indent_first_line))

    return elem


def table_style_to_xml(style: TableStyleInfo) -> ET.Element:
    """Convert a TableStyleInfo to an XML element.

    Args:
        style: The TableStyleInfo to convert.

    Returns:
        XML Element representing the table style.
    """
    elem = ET.Element('table-style')
    elem.set('name', style.name)

    if style.header_fill_color:
        header = ET.SubElement(elem, 'header')
        header.set('fill', style.header_fill_color)
        if style.header_text_color:
            header.set('text-color', style.header_text_color)
        header.set('bold', 'true' if style.header_bold else 'false')

    if style.row_band_color:
        band = ET.SubElement(elem, 'row-banding')
        band.set('color', style.row_band_color)

    if style.border_color or style.border_width:
        border = ET.SubElement(elem, 'border')
        if style.border_color:
            border.set('color', style.border_color)
        if style.border_width:
            border.set('width', str(style.border_width))

    return elem


def save_styles(
    styles: Dict[str, StyleInfo],
    table_styles: List[TableStyleInfo],
    output_path: Path
) -> None:
    """Save extracted styles to an XML file.

    Args:
        styles: Dictionary of paragraph styles.
        table_styles: List of table styles.
        output_path: Path to save the styles XML.
    """
    root = ET.Element('styles')
    root.set('version', '1.0')

    # Add paragraph styles
    para_styles = ET.SubElement(root, 'paragraph-styles')
    for style in styles.values():
        para_styles.append(style_to_xml(style))

    # Add table styles
    tbl_styles = ET.SubElement(root, 'table-styles')
    for style in table_styles:
        tbl_styles.append(table_style_to_xml(style))

    # Pretty print and save
    rough_string = ET.tostring(root, encoding='unicode')
    reparsed = minidom.parseString(rough_string)
    formatted = reparsed.toprettyxml(indent="  ")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(formatted, encoding='utf-8')


def load_styles(styles_path: Path) -> Dict[str, StyleInfo]:
    """Load styles from an XML file.

    Args:
        styles_path: Path to the styles XML file.

    Returns:
        Dictionary mapping style names to StyleInfo objects.
    """
    if not styles_path.exists():
        raise FileNotFoundError(f"Styles file not found: {styles_path}")

    tree = ET.parse(styles_path)
    root = tree.getroot()

    styles = {}

    para_styles = root.find('paragraph-styles')
    if para_styles is not None:
        for style_elem in para_styles.findall('style'):
            name = style_elem.get('name', '')
            info = StyleInfo(name=name)

            # Parse font
            font = style_elem.find('font')
            if font is not None:
                info.font_name = font.get('name')
                size = font.get('size')
                if size:
                    info.font_size = int(size)
                info.bold = font.get('bold') == 'true'
                info.italic = font.get('italic') == 'true'
                info.underline = font.get('underline') == 'true'

            # Parse color
            color = style_elem.find('color')
            if color is not None:
                info.color = color.get('value')

            # Parse fill
            fill = style_elem.find('fill')
            if fill is not None:
                info.background_color = fill.get('color')

            # Parse spacing
            spacing = style_elem.find('spacing')
            if spacing is not None:
                before = spacing.get('before')
                after = spacing.get('after')
                line = spacing.get('line')
                if before:
                    info.spacing_before = int(before)
                if after:
                    info.spacing_after = int(after)
                if line:
                    info.line_spacing = float(line)

            # Parse alignment
            align = style_elem.find('alignment')
            if align is not None:
                info.alignment = align.get('value')

            # Parse indent
            indent = style_elem.find('indent')
            if indent is not None:
                left = indent.get('left')
                right = indent.get('right')
                first = indent.get('first-line')
                if left:
                    info.indent_left = int(left)
                if right:
                    info.indent_right = int(right)
                if first:
                    info.indent_first_line = int(first)

            styles[name] = info

    return styles


def apply_styles(doc: Document, styles: Dict[str, StyleInfo]) -> None:
    """Apply loaded styles to a document.

    This function modifies the document in place to apply the given styles.

    Args:
        doc: The Document to modify.
        styles: Dictionary of styles to apply.
    """
    # This is a placeholder for style application logic
    # Full implementation would require modifying the document's style definitions
    # which is complex with python-docx
    pass


def get_styles_path(document_path: Path) -> Path:
    """Get the styles file path for a document.

    Args:
        document_path: Path to the Markdown or DOCX file.

    Returns:
        Path to the corresponding _styles.xml file.
    """
    return document_path.parent / f"{document_path.stem}_styles.xml"


def save_styles_to_both_locations(
    styles: Dict[str, StyleInfo],
    table_styles: List[TableStyleInfo],
    source_path: Path,
    output_path: Path
) -> None:
    """Save extracted styles to both source and output locations.

    Args:
        styles: Dictionary of paragraph styles.
        table_styles: List of table styles.
        source_path: Path to the source file (DOCX).
        output_path: Path to the output file (MD).
    """
    # Get styles paths for both locations
    source_styles = get_styles_path(source_path)
    output_styles = get_styles_path(output_path)

    # Save to source location
    save_styles(styles, table_styles, source_styles)

    # Save to output location if different
    if source_styles != output_styles:
        save_styles(styles, table_styles, output_styles)
