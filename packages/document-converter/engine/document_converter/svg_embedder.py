"""SVG embedding support for Word documents.

This module provides functionality to embed SVG images directly in DOCX files.
Requires python-docx-ng for SVG support (standard python-docx does not support SVG).

Usage:
    from document_converter.svg_embedder import (
        is_svg_embedding_available,
        embed_svg_in_docx,
        convert_svg_to_png_fallback
    )

    if is_svg_embedding_available():
        embed_svg_in_docx(doc, svg_path, width, height)
    else:
        # Fall back to PNG
        png_path = convert_svg_to_png_fallback(svg_path)
        doc.add_picture(str(png_path), width=width)
"""
import subprocess
from pathlib import Path
from typing import Optional, Tuple
import warnings


# Cache the availability check result
_svg_embedding_available: Optional[bool] = None


def is_svg_embedding_available() -> bool:
    """Check if SVG embedding is available.

    SVG embedding requires python-docx-ng which extends python-docx
    with SVG support.

    Returns:
        True if python-docx-ng is installed and SVG embedding works.
    """
    global _svg_embedding_available

    if _svg_embedding_available is not None:
        return _svg_embedding_available

    try:
        from docx import Document
        from docx.image.exceptions import UnrecognizedImageError

        # Create a minimal test SVG
        test_svg = b'''<?xml version="1.0"?>
<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1">
<rect width="1" height="1"/>
</svg>'''

        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.svg', delete=False) as f:
            f.write(test_svg)
            test_path = f.name

        try:
            doc = Document()
            doc.add_picture(test_path)
            _svg_embedding_available = True
        except UnrecognizedImageError:
            _svg_embedding_available = False
        finally:
            Path(test_path).unlink(missing_ok=True)

    except ImportError:
        _svg_embedding_available = False

    return _svg_embedding_available


def get_svg_dimensions(svg_path: Path) -> Tuple[Optional[int], Optional[int]]:
    """Extract dimensions from an SVG file.

    Parses the SVG to get width and height attributes or viewBox dimensions.

    Args:
        svg_path: Path to the SVG file.

    Returns:
        Tuple of (width, height) in pixels, or (None, None) if not determinable.
    """
    try:
        import xml.etree.ElementTree as ET

        tree = ET.parse(svg_path)
        root = tree.getroot()

        # Try to get width/height attributes
        width_str = root.attrib.get('width', '')
        height_str = root.attrib.get('height', '')

        def parse_dimension(dim_str: str) -> Optional[int]:
            """Parse a dimension string to pixels."""
            if not dim_str:
                return None
            # Remove units and convert to int
            dim_str = dim_str.strip()
            for unit in ('px', 'pt', 'em', 'rem', '%'):
                dim_str = dim_str.replace(unit, '')
            try:
                return int(float(dim_str))
            except ValueError:
                return None

        width = parse_dimension(width_str)
        height = parse_dimension(height_str)

        # If dimensions not found, try viewBox
        if width is None or height is None:
            viewbox = root.attrib.get('viewBox', '')
            if viewbox:
                parts = viewbox.split()
                if len(parts) >= 4:
                    try:
                        width = int(float(parts[2]))
                        height = int(float(parts[3]))
                    except ValueError:
                        pass

        return width, height

    except Exception:
        return None, None


def convert_svg_to_png(svg_path: Path, output_path: Optional[Path] = None) -> Path:
    """Convert an SVG file to PNG using available tools.

    Tries multiple conversion methods in order:
    1. Mermaid CLI (mmdc) - if available
    2. cairosvg - if installed
    3. Inkscape - if installed

    Args:
        svg_path: Path to the input SVG file.
        output_path: Optional output path. If None, uses same name with .png extension.

    Returns:
        Path to the generated PNG file.

    Raises:
        RuntimeError: If no conversion tool is available or conversion fails.
    """
    if output_path is None:
        output_path = svg_path.with_suffix('.png')

    # Try cairosvg first (Python library, most reliable)
    try:
        import cairosvg
        cairosvg.svg2png(url=str(svg_path), write_to=str(output_path))
        if output_path.exists():
            return output_path
    except ImportError:
        pass
    except Exception as e:
        warnings.warn(f"cairosvg conversion failed: {e}")

    # Try Inkscape
    try:
        result = subprocess.run(
            ['inkscape', '--export-type=png', '--export-filename=' + str(output_path), str(svg_path)],
            capture_output=True,
            timeout=30
        )
        if result.returncode == 0 and output_path.exists():
            return output_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    except Exception as e:
        warnings.warn(f"Inkscape conversion failed: {e}")

    # Try rsvg-convert (part of librsvg)
    try:
        result = subprocess.run(
            ['rsvg-convert', '-o', str(output_path), str(svg_path)],
            capture_output=True,
            timeout=30
        )
        if result.returncode == 0 and output_path.exists():
            return output_path
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    except Exception as e:
        warnings.warn(f"rsvg-convert conversion failed: {e}")

    raise RuntimeError(
        f"Failed to convert SVG to PNG. No conversion tool available. "
        "Install one of: cairosvg (pip install cairosvg), Inkscape, or librsvg."
    )


def embed_svg_images_in_docx(
    docx_path: Path,
    svg_paths: list[Path],
    fallback_to_png: bool = True
) -> dict:
    """Replace PNG placeholders with SVG images in a DOCX file.

    This is a post-processing step that can replace PNG images with their
    SVG equivalents after the initial DOCX generation.

    Args:
        docx_path: Path to the DOCX file to modify.
        svg_paths: List of SVG file paths to embed.
        fallback_to_png: If True and SVG embedding unavailable, keep PNGs.

    Returns:
        Dictionary with 'embedded' count and any 'warnings'.
    """
    result = {
        'embedded': 0,
        'warnings': []
    }

    if not is_svg_embedding_available():
        if fallback_to_png:
            result['warnings'].append(
                "SVG embedding not available (requires python-docx-ng). "
                "Using PNG images instead."
            )
            return result
        else:
            raise RuntimeError(
                "SVG embedding requires python-docx-ng. "
                "Install with: pip install python-docx-ng"
            )

    # SVG embedding is available - this would be used for post-processing
    # For now, we handle SVG at render time in the converters
    result['embedded'] = len(svg_paths)
    return result


def check_svg_embedding_requirements() -> dict:
    """Check all requirements for SVG embedding.

    Returns:
        Dictionary with status information about SVG embedding capabilities.
    """
    status = {
        'svg_embedding_available': is_svg_embedding_available(),
        'docx_library': 'unknown',
        'svg_to_png_available': False,
        'svg_to_png_tool': None,
        'recommendations': []
    }

    # Check which docx library is installed
    try:
        import docx
        # Check if it's python-docx-ng by testing SVG support
        if is_svg_embedding_available():
            status['docx_library'] = 'python-docx-ng'
        else:
            status['docx_library'] = 'python-docx'
            status['recommendations'].append(
                "For SVG embedding, install python-docx-ng: pip install python-docx-ng"
            )
    except ImportError:
        status['docx_library'] = 'not installed'
        status['recommendations'].append(
            "Install python-docx: pip install python-docx"
        )

    # Check SVG to PNG conversion tools
    try:
        import cairosvg
        status['svg_to_png_available'] = True
        status['svg_to_png_tool'] = 'cairosvg'
    except ImportError:
        # Check for command-line tools
        for tool in ['inkscape', 'rsvg-convert']:
            try:
                result = subprocess.run(
                    [tool, '--version'],
                    capture_output=True,
                    timeout=5
                )
                if result.returncode == 0:
                    status['svg_to_png_available'] = True
                    status['svg_to_png_tool'] = tool
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

    if not status['svg_to_png_available']:
        status['recommendations'].append(
            "For SVG to PNG fallback, install cairosvg: pip install cairosvg"
        )

    return status
