"""Table formatting utilities for DOCX documents."""
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Twips
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.table import Table


@dataclass
class TableStyleConfig:
    """Configuration for table styling.

    Attributes:
        header_bg_color: Background color for header row (hex without #).
        header_text_color: Text color for header row (hex without #).
        odd_row_color: Background color for odd data rows (hex without #).
        even_row_color: Background color for even data rows (hex without #).
        border_color: Border color (hex without #).
        border_size: Border size in eighths of a point.
        striped: Whether to use alternating row colors.
        column_widths: Optional dict mapping column names to width multipliers.
    """
    header_bg_color: str = "4472C4"
    header_text_color: str = "FFFFFF"
    odd_row_color: str = "FFFFFF"
    even_row_color: str = "D6DCE5"
    border_color: str = "8EAADB"
    border_size: int = 4
    striped: bool = True
    column_widths: Optional[Dict[str, float]] = None

    @classmethod
    def from_dict(cls, style_dict: Dict[str, Any]) -> 'TableStyleConfig':
        """Create TableStyleConfig from a dictionary.

        Args:
            style_dict: Dictionary with style properties.

        Returns:
            TableStyleConfig with values from dict, defaults for missing keys.
        """
        return cls(
            header_bg_color=style_dict.get('header_bg_color', "4472C4"),
            header_text_color=style_dict.get('header_text_color', "FFFFFF"),
            odd_row_color=style_dict.get('odd_row_color', "FFFFFF"),
            even_row_color=style_dict.get('even_row_color', "D6DCE5"),
            border_color=style_dict.get('border_color', "8EAADB"),
            border_size=style_dict.get('border_size', 4),
            striped=style_dict.get('striped', True),
            column_widths=style_dict.get('column_widths'),
        )

    def merge(self, override: 'TableStyleConfig') -> 'TableStyleConfig':
        """Merge this config with an override config.

        Override values take precedence where they differ from defaults.

        Args:
            override: TableStyleConfig with override values.

        Returns:
            New TableStyleConfig with merged values.
        """
        default = TableStyleConfig()
        return TableStyleConfig(
            header_bg_color=override.header_bg_color if override.header_bg_color != default.header_bg_color else self.header_bg_color,
            header_text_color=override.header_text_color if override.header_text_color != default.header_text_color else self.header_text_color,
            odd_row_color=override.odd_row_color if override.odd_row_color != default.odd_row_color else self.odd_row_color,
            even_row_color=override.even_row_color if override.even_row_color != default.even_row_color else self.even_row_color,
            border_color=override.border_color if override.border_color != default.border_color else self.border_color,
            border_size=override.border_size if override.border_size != default.border_size else self.border_size,
            striped=override.striped if override.striped != default.striped else self.striped,
            column_widths=override.column_widths if override.column_widths is not None else self.column_widths,
        )


# Default table style configuration
DEFAULT_TABLE_STYLE = TableStyleConfig()


def parse_table_style_block(content: str) -> List[Tuple[TableStyleConfig, int]]:
    """Parse table style blocks from content.

    Finds HTML comments in format: <!-- table-style: key=value, key=value -->

    Args:
        content: The document content to parse.

    Returns:
        List of tuples (TableStyleConfig, position) for each style block found.
    """
    blocks = []

    # Pattern to match table-style comments
    pattern = r'<!--\s*table-style:\s*(.+?)\s*-->'

    for match in re.finditer(pattern, content):
        style_str = match.group(1)
        position = match.start()

        # Parse key=value pairs
        style_dict: Dict[str, Any] = {}

        # Handle column_widths JSON specially
        col_widths_match = re.search(r'column_widths=(\{[^}]+\})', style_str)
        if col_widths_match:
            try:
                style_dict['column_widths'] = json.loads(col_widths_match.group(1))
            except json.JSONDecodeError:
                pass
            # Remove from style_str for further parsing
            style_str = style_str.replace(col_widths_match.group(0), '')

        # Parse remaining key=value pairs
        # Keys that should remain as strings (hex colors)
        color_keys = {'header_bg_color', 'header_text_color', 'odd_row_color',
                      'even_row_color', 'border_color'}

        for pair in re.findall(r'(\w+)=([^,\s]+)', style_str):
            key, value = pair
            # Convert types
            if value.lower() == 'true':
                style_dict[key] = True
            elif value.lower() == 'false':
                style_dict[key] = False
            elif key in color_keys:
                # Keep color values as strings
                style_dict[key] = value
            elif value.isdigit():
                style_dict[key] = int(value)
            else:
                style_dict[key] = value

        config = TableStyleConfig.from_dict(style_dict)
        blocks.append((config, position))

    return blocks


def apply_column_widths(table: Table, widths: Dict[str, float]) -> None:
    """Apply column width multipliers to a table.

    Args:
        table: The docx Table to modify.
        widths: Dict mapping column header text to width multipliers.
    """
    if not table.rows:
        return

    # Get header row to map column names to indices
    header_row = table.rows[0]
    col_indices: Dict[str, int] = {}

    for idx, cell in enumerate(header_row.cells):
        cell_text = cell.text.strip()
        if cell_text in widths:
            col_indices[cell_text] = idx

    if not col_indices:
        return

    # Calculate total width multiplier for normalization
    total_mult = sum(widths.get(cell.text.strip(), 1.0) for cell in header_row.cells)

    # Get current table width or use default
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    table_width = 9360  # Default ~6.5 inches in twips

    if tbl_pr is not None:
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is not None:
            width_val = tbl_w.get(qn('w:w'))
            if width_val and width_val.isdigit():
                table_width = int(width_val)

    # Apply widths to each column
    for cell in header_row.cells:
        cell_text = cell.text.strip()
        mult = widths.get(cell_text, 1.0)
        col_width = int(table_width * mult / total_mult)

        # Set cell width
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{col_width}" w:type="dxa"/>')
            tcPr.append(tcW)
        else:
            tcW.set(qn('w:w'), str(col_width))
            tcW.set(qn('w:type'), 'dxa')


def get_table_headers(table: Table) -> List[str]:
    """Get the header row column names from a table.

    Args:
        table: The docx Table to inspect.

    Returns:
        List of column header texts, or empty list if no rows.
    """
    if not table.rows:
        return []
    return [cell.text.strip() for cell in table.rows[0].cells]


def find_matching_profile(
    table: Table,
    profiles: Dict[str, Dict[str, float]]
) -> Optional[Dict[str, float]]:
    """Find a matching column width profile for a table.

    Matches profiles by checking if the table's headers match the profile's
    column names exactly (same columns in any order).

    Args:
        table: The docx Table to match.
        profiles: Dict mapping profile names to column width dicts.

    Returns:
        The matching profile's column widths dict, or None if no match.
    """
    if not profiles:
        return None

    headers = get_table_headers(table)
    if not headers:
        return None

    header_set = set(headers)

    for profile_name, profile_widths in profiles.items():
        profile_columns = set(profile_widths.keys())
        if header_set == profile_columns:
            return profile_widths

    return None


def apply_table_profiles(
    docx_path: Path,
    profiles: Dict[str, Dict[str, float]]
) -> int:
    """Apply column width profiles to all matching tables in a document.

    For each table, checks if its headers match any profile and applies
    the corresponding column widths.

    Args:
        docx_path: Path to the DOCX file.
        profiles: Dict mapping profile names to column width dicts.

    Returns:
        Number of tables that had profiles applied.
    """
    if not profiles:
        return 0

    doc = Document(docx_path)
    applied_count = 0

    for table in doc.tables:
        matching_widths = find_matching_profile(table, profiles)
        if matching_widths:
            apply_column_widths(table, matching_widths)
            applied_count += 1

    if applied_count > 0:
        doc.save(docx_path)

    return applied_count


def set_cell_shading(cell, color: str) -> None:
    """Set background shading color for a cell.

    Args:
        cell: The table cell to shade.
        color: Hex color code without # (e.g., "4472C4" for blue).
    """
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, color: str = "000000", size: int = 4) -> None:
    """Set borders for a cell.

    Args:
        cell: The table cell.
        color: Border color in hex without # (default: black).
        size: Border size in eighths of a point (default: 4 = 0.5pt).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = parse_xml(
        f'''<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{size}" w:color="{color}"/>
            <w:left w:val="single" w:sz="{size}" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>
            <w:right w:val="single" w:sz="{size}" w:color="{color}"/>
        </w:tcBorders>'''
    )
    tcPr.append(tcBorders)


def format_table(
    table: Table,
    header_bg_color: str = "4472C4",
    header_text_color: str = "FFFFFF",
    odd_row_color: str = "FFFFFF",
    even_row_color: str = "D6DCE5",
    border_color: str = "8EAADB",
    border_size: int = 4
) -> None:
    """Format a single table with borders, header styling, and striped rows.

    Args:
        table: The docx Table object to format.
        header_bg_color: Background color for header row (hex without #).
        header_text_color: Text color for header row (hex without #).
        odd_row_color: Background color for odd rows (hex without #).
        even_row_color: Background color for even rows (hex without #).
        border_color: Border color (hex without #).
        border_size: Border size in eighths of a point.
    """
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            # Set borders on all cells
            set_cell_borders(cell, color=border_color, size=border_size)

            if row_idx == 0:
                # Header row: dark background, white text
                set_cell_shading(cell, header_bg_color)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(
                            int(header_text_color[0:2], 16),
                            int(header_text_color[2:4], 16),
                            int(header_text_color[4:6], 16)
                        )
                        run.font.bold = True
            else:
                # Data rows: alternating colors
                if row_idx % 2 == 1:
                    set_cell_shading(cell, odd_row_color)
                else:
                    set_cell_shading(cell, even_row_color)


def format_tables(
    docx_path: Path,
    header_bg_color: str = "4472C4",
    header_text_color: str = "FFFFFF",
    odd_row_color: str = "FFFFFF",
    even_row_color: str = "D6DCE5",
    border_color: str = "8EAADB",
    border_size: int = 4
) -> None:
    """Format all tables in a DOCX document.

    Args:
        docx_path: Path to the DOCX file to format.
        header_bg_color: Background color for header rows (hex without #).
        header_text_color: Text color for header rows (hex without #).
        odd_row_color: Background color for odd data rows (hex without #).
        even_row_color: Background color for even data rows (hex without #).
        border_color: Border color (hex without #).
        border_size: Border size in eighths of a point.
    """
    doc = Document(docx_path)

    for table in doc.tables:
        format_table(
            table,
            header_bg_color=header_bg_color,
            header_text_color=header_text_color,
            odd_row_color=odd_row_color,
            even_row_color=even_row_color,
            border_color=border_color,
            border_size=border_size
        )

    doc.save(docx_path)


def fix_table_styles_for_pdf(docx_path: Path) -> dict:
    """Fix table styles that cause PDF rendering issues in LibreOffice.

    Pandoc generates tables with invalid style references ('Table', 'Compact')
    that don't exist in the document's style definitions. LibreOffice fails to
    render these tables correctly, showing only the last cell content.

    This function removes invalid style references while preserving the actual
    formatting (borders, shading, etc.) that was applied directly to elements.

    Args:
        docx_path: Path to the DOCX file to fix.

    Returns:
        Dictionary with counts of fixes applied.
    """
    doc = Document(docx_path)
    fixes = {'tblStyle': 0, 'pStyle_Compact': 0}

    for table in doc.tables:
        tbl_elem = table._tbl
        tbl_pr = tbl_elem.find(qn('w:tblPr'))

        if tbl_pr is not None:
            # Remove invalid 'Table' style reference
            # The actual formatting (borders, shading) is applied directly to cells
            tbl_style = tbl_pr.find(qn('w:tblStyle'))
            if tbl_style is not None:
                style_val = tbl_style.get(qn('w:val'))
                if style_val == 'Table':
                    tbl_pr.remove(tbl_style)
                    fixes['tblStyle'] += 1

        # Remove invalid 'Compact' paragraph style from table cells
        # Add proper spacing to maintain compact appearance
        for p in tbl_elem.findall('.//' + qn('w:p')):
            pPr = p.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None and pStyle.get(qn('w:val')) == 'Compact':
                    pPr.remove(pStyle)
                    fixes['pStyle_Compact'] += 1

                    # Add spacing to maintain compact look
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = parse_xml(
                            '<w:spacing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                            'w:before="40" w:after="40" w:line="240" w:lineRule="auto"/>'
                        )
                        pPr.append(spacing)

    doc.save(docx_path)
    return fixes
