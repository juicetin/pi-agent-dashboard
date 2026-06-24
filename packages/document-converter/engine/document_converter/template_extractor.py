"""Template extraction module for creating reusable DOCX templates."""
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
from datetime import datetime
from xml.etree import ElementTree as ET
from xml.dom import minidom
import shutil
import re

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from .style_extractor import extract_styles, extract_table_styles, save_styles
from .textbox_extractor import (
    extract_textboxes_from_docx,
    display_textboxes,
    replace_text_in_textbox,
    TextBoxElement
)
from .font_manager import (
    extract_fonts_from_docx,
    extract_embedded_fonts,
    find_font_file,
    get_font_name_from_file,
)


@dataclass
class DocumentElement:
    """Represents an element that can be converted to a variable."""
    index: int
    content: str
    location: str  # 'cover', 'header', 'footer', 'body'
    position: str  # 'title', 'author', 'date', 'left', 'center', 'right', etc.
    style_info: str  # Human-readable style description
    element_path: str  # Path to element in document
    paragraph_index: int = 0


@dataclass
class TemplateVariable:
    """A variable defined in the template."""
    name: str
    location: str
    position: str
    original_content: str
    syntax: str = "both"  # 'angle', 'brace', or 'both'


def get_style_description(paragraph) -> str:
    """Get a human-readable style description for a paragraph.

    Args:
        paragraph: The paragraph to describe.

    Returns:
        String describing the style (e.g., "36pt bold").
    """
    parts = []

    # Get font size
    for run in paragraph.runs:
        if run.font.size:
            parts.append(f"{run.font.size.pt}pt")
            break

    # Check for bold
    for run in paragraph.runs:
        if run.bold:
            parts.append("bold")
            break

    # Check for italic
    for run in paragraph.runs:
        if run.italic:
            parts.append("italic")
            break

    # Get style name
    if paragraph.style and paragraph.style.name:
        parts.append(f"({paragraph.style.name})")

    return ' '.join(parts) if parts else "normal"


def has_page_break(paragraph) -> bool:
    """Check if a paragraph contains a page break.

    Args:
        paragraph: The paragraph to check.

    Returns:
        True if the paragraph contains a page break.
    """
    for run in paragraph.runs:
        for elem in run._r:
            if elem.tag == qn('w:br'):
                br_type = elem.get(qn('w:type'))
                if br_type == 'page':
                    return True
    return False


def display_paragraphs_for_selection(doc: Document, max_paragraphs: int = 30) -> List[dict]:
    """Display numbered paragraphs with styles and page breaks for user selection.

    Args:
        doc: The document to analyze.
        max_paragraphs: Maximum number of paragraphs to display.

    Returns:
        List of paragraph info dictionaries.
    """
    paragraphs_info = []

    print("\n--- Document Paragraphs ---")
    print("(Use paragraph number or 'p' for first page break)")
    print()

    for i, para in enumerate(doc.paragraphs[:max_paragraphs]):
        text = para.text.strip()
        style_name = para.style.name if para.style else "Normal"
        style_desc = get_style_description(para)

        # Check for page break
        is_page_break = has_page_break(para)

        info = {
            'index': i + 1,
            'text': text[:60] + ('...' if len(text) > 60 else ''),
            'style': style_name,
            'style_desc': style_desc,
            'has_page_break': is_page_break
        }
        paragraphs_info.append(info)

        # Display
        page_break_marker = " [PAGE BREAK]" if is_page_break else ""
        if text:
            print(f"  [{i + 1}] \"{info['text']}\"")
            print(f"       Style: {style_name} ({style_desc}){page_break_marker}")
        elif is_page_break:
            print(f"  [{i + 1}] <empty paragraph> [PAGE BREAK]")
        else:
            print(f"  [{i + 1}] <empty paragraph>")

    return paragraphs_info


def prompt_cover_boundary(doc: Document) -> int:
    """Prompt user to specify the cover page boundary.

    Args:
        doc: The document to analyze.

    Returns:
        Paragraph index (0-based) where cover page ends.
        -1 if no cover page should be kept.
    """
    paragraphs_info = display_paragraphs_for_selection(doc)

    print("\nEnter cover page boundary:")
    print("  - Paragraph number (e.g., '5') = keep paragraphs 1-5 as cover page")
    print("  - 'p' = use first page break as boundary")
    print("  - 'none' or 0 = no cover page")
    print()

    try:
        response = input("Cover page ends after: ").strip().lower()
    except EOFError:
        return -1

    if response == 'none' or response == '0' or response == '':
        return -1

    if response == 'p':
        # Find first page break
        for info in paragraphs_info:
            if info['has_page_break']:
                print(f"Using page break at paragraph {info['index']} as boundary.")
                return info['index'] - 1  # Return 0-based index
        print("No page break found. Using first 5 paragraphs as cover.")
        return 4  # Default to first 5 paragraphs

    try:
        para_num = int(response)
        if 1 <= para_num <= len(paragraphs_info):
            return para_num - 1  # Convert to 0-based index
        else:
            print(f"Invalid paragraph number. Using default (5).")
            return 4
    except ValueError:
        print("Invalid input. Using default (5).")
        return 4


def find_cover_page_elements(
    doc: Document,
    cover_boundary: int = -1
) -> List[DocumentElement]:
    """Find potential variable elements on the cover page.

    Args:
        doc: The document to analyze.
        cover_boundary: Paragraph index (0-based) where cover page ends.
                       If -1, auto-detect based on page breaks or headings.

    Returns:
        List of DocumentElement objects representing cover page elements.
    """
    elements = []
    element_index = 1

    # Determine the range of paragraphs to consider
    if cover_boundary >= 0:
        max_para = cover_boundary + 1
    else:
        max_para = 20  # Default: check first 20 paragraphs

    # Analyze paragraphs within the cover page boundary
    for i, para in enumerate(doc.paragraphs[:max_para]):
        text = para.text.strip()
        if not text:
            continue

        # If no explicit boundary, stop at regular headings
        if cover_boundary < 0:
            style_name = para.style.name if para.style else ""
            if style_name.startswith('Heading') and style_name != 'Heading 1':
                break

        # Determine position based on style and content
        style_name = para.style.name if para.style else ""
        position = "unknown"
        if style_name == 'Title' or (i == 0 and len(text) < 100):
            position = "title"
        elif style_name == 'Subtitle':
            position = "subtitle"
        elif _looks_like_date(text):
            position = "date"
        elif _looks_like_author(text):
            position = "author"
        elif len(text) < 50:
            position = f"line_{i + 1}"

        style_desc = get_style_description(para)

        elements.append(DocumentElement(
            index=element_index,
            content=text[:50] + ('...' if len(text) > 50 else ''),
            location='cover',
            position=position,
            style_info=style_desc,
            element_path=f"/w:body/w:p[{i + 1}]",
            paragraph_index=i
        ))
        element_index += 1

        # If no explicit boundary, stop after finding several elements
        if cover_boundary < 0 and element_index > 10:
            break

    return elements


def _looks_like_date(text: str) -> bool:
    """Check if text looks like a date."""
    date_patterns = [
        r'\d{4}[-/]\d{1,2}[-/]\d{1,2}',  # 2024-01-15
        r'\d{1,2}[-/]\d{1,2}[-/]\d{4}',  # 15/01/2024
        r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}',
        r'\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
    ]
    for pattern in date_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return True
    return False


def _looks_like_author(text: str) -> bool:
    """Check if text looks like an author name."""
    # Simple heuristic: 2-4 words, title case, no numbers
    if re.search(r'\d', text):
        return False
    words = text.split()
    if 2 <= len(words) <= 4:
        if all(word[0].isupper() for word in words if word):
            return True
    return False


def find_header_elements(doc: Document) -> List[DocumentElement]:
    """Find elements in document headers (default header only).

    Args:
        doc: The document to analyze.

    Returns:
        List of DocumentElement objects representing header elements.
    """
    elements = []
    element_index = 1

    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            continue

        for i, para in enumerate(header.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Determine position (left/center/right) based on alignment
            alignment = para.alignment
            if alignment == 1:  # Center
                position = "center"
            elif alignment == 2:  # Right
                position = "right"
            else:
                position = "left"

            style_desc = get_style_description(para)

            elements.append(DocumentElement(
                index=element_index,
                content=text[:50] + ('...' if len(text) > 50 else ''),
                location='header',
                position=position,
                style_info=style_desc,
                element_path=f"/w:sectPr/w:headerReference/w:p[{i + 1}]",
                paragraph_index=i
            ))
            element_index += 1

    return elements


def find_first_page_header_elements(doc: Document) -> List[DocumentElement]:
    """Find elements in the first page (cover) header.

    Args:
        doc: The document to analyze.

    Returns:
        List of DocumentElement objects representing first page header elements.
    """
    elements = []
    element_index = 1

    for section in doc.sections:
        # Check if section has different first page header
        if not section.different_first_page_header_footer:
            continue

        first_header = section.first_page_header
        if not first_header or not first_header.paragraphs:
            continue

        for i, para in enumerate(first_header.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            alignment = para.alignment
            if alignment == 1:
                position = "center"
            elif alignment == 2:
                position = "right"
            else:
                position = "left"

            style_desc = get_style_description(para)

            elements.append(DocumentElement(
                index=element_index,
                content=text[:50] + ('...' if len(text) > 50 else ''),
                location='header_first',
                position=position,
                style_info=style_desc,
                element_path=f"/w:sectPr/w:headerReference[@w:type='first']/w:p[{i + 1}]",
                paragraph_index=i
            ))
            element_index += 1

    return elements


def find_footer_elements(doc: Document) -> List[DocumentElement]:
    """Find elements in document footers (default footer only).

    Args:
        doc: The document to analyze.

    Returns:
        List of DocumentElement objects representing footer elements.
    """
    elements = []
    element_index = 1

    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            continue

        for i, para in enumerate(footer.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # Determine position
            alignment = para.alignment
            if alignment == 1:
                position = "center"
            elif alignment == 2:
                position = "right"
            else:
                position = "left"

            style_desc = get_style_description(para)

            elements.append(DocumentElement(
                index=element_index,
                content=text[:50] + ('...' if len(text) > 50 else ''),
                location='footer',
                position=position,
                style_info=style_desc,
                element_path=f"/w:sectPr/w:footerReference/w:p[{i + 1}]",
                paragraph_index=i
            ))
            element_index += 1

    return elements


def find_first_page_footer_elements(doc: Document) -> List[DocumentElement]:
    """Find elements in the first page (cover) footer.

    Args:
        doc: The document to analyze.

    Returns:
        List of DocumentElement objects representing first page footer elements.
    """
    elements = []
    element_index = 1

    for section in doc.sections:
        if not section.different_first_page_header_footer:
            continue

        first_footer = section.first_page_footer
        if not first_footer or not first_footer.paragraphs:
            continue

        for i, para in enumerate(first_footer.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            alignment = para.alignment
            if alignment == 1:
                position = "center"
            elif alignment == 2:
                position = "right"
            else:
                position = "left"

            style_desc = get_style_description(para)

            elements.append(DocumentElement(
                index=element_index,
                content=text[:50] + ('...' if len(text) > 50 else ''),
                location='footer_first',
                position=position,
                style_info=style_desc,
                element_path=f"/w:sectPr/w:footerReference[@w:type='first']/w:p[{i + 1}]",
                paragraph_index=i
            ))
            element_index += 1

    return elements


def present_elements(elements: List[DocumentElement], section_name: str) -> str:
    """Format elements for display to the user.

    Args:
        elements: List of elements to display.
        section_name: Name of the section (e.g., "cover page").

    Returns:
        Formatted string for display.
    """
    if not elements:
        return f"No {section_name} elements found.\n"

    lines = [f"\nFound {section_name} elements:"]
    for elem in elements:
        lines.append(f'  [{elem.index}] "{elem.content}" ({elem.style_info})')

    return '\n'.join(lines)


def prompt_variable_selection(elements: List[DocumentElement], section_name: str) -> List[DocumentElement]:
    """Prompt user to select which elements should become variables.

    Args:
        elements: List of available elements.
        section_name: Name of the section for the prompt.

    Returns:
        List of selected elements.
    """
    if not elements:
        return []

    print(present_elements(elements, section_name))
    print(f"\nWhich {section_name} elements should be substitutable? (comma-separated, or 'none'): ", end='')

    try:
        response = input().strip().lower()
    except EOFError:
        return []

    if response == 'none' or response == '':
        return []

    selected = []
    try:
        indices = [int(x.strip()) for x in response.split(',')]
        for idx in indices:
            for elem in elements:
                if elem.index == idx:
                    selected.append(elem)
                    break
    except ValueError:
        print("Invalid input. No elements selected.")
        return []

    return selected


def prompt_textbox_variable_selection(
    textboxes: List[TextBoxElement]
) -> Tuple[List[Tuple[TextBoxElement, int, str]], Dict[str, str]]:
    """Prompt user to select text box lines as variables.

    Args:
        textboxes: List of text box elements.

    Returns:
        Tuple of (selected items as (textbox, line_index, var_name), replacements dict).
    """
    if not textboxes:
        return [], {}

    display_textboxes(textboxes)

    print("\nWhich text box lines should be variables?")
    print("Format: textbox.line (e.g., '1.1,1.2,1.3' for lines 1-3 of textbox 1)")
    print("Or 'none' to skip: ", end='')

    try:
        response = input().strip().lower()
    except EOFError:
        return [], {}

    if response == 'none' or response == '':
        return [], {}

    selected = []
    replacements = {}

    for item in response.split(','):
        item = item.strip()
        if '.' in item:
            try:
                tb_idx, line_idx = item.split('.')
                tb_idx = int(tb_idx)
                line_idx = int(line_idx)

                for tb in textboxes:
                    if tb.index == tb_idx and line_idx <= len(tb.paragraphs):
                        para_text = tb.paragraphs[line_idx - 1]
                        display = para_text[:50] + ('...' if len(para_text) > 50 else '')

                        print(f"\nVariable name for \"{display}\": ", end='')
                        try:
                            var_name = input().strip()
                        except EOFError:
                            var_name = f"textbox_{tb_idx}_{line_idx}"

                        if not var_name:
                            var_name = f"textbox_{tb_idx}_{line_idx}"

                        var_name = re.sub(r'[^a-zA-Z0-9_]', '_', var_name)
                        replacements[para_text] = f"{{{{{var_name}}}}}"
                        selected.append((tb, line_idx, var_name, para_text))
                        break
            except ValueError:
                continue

    return selected, replacements


def textbox_selections_to_variables(
    selections: List[Tuple[TextBoxElement, int, str, str]]
) -> List[TemplateVariable]:
    """Convert text box selections to TemplateVariable objects.

    Args:
        selections: List of (textbox, line_index, var_name, original_text) tuples.

    Returns:
        List of TemplateVariable objects.
    """
    variables = []
    for tb, line_idx, var_name, original_text in selections:
        variables.append(TemplateVariable(
            name=var_name,
            location='textbox',
            position=f"line_{line_idx}",
            original_content=original_text,
            syntax='both'
        ))
    return variables


def prompt_variable_names(elements: List[DocumentElement]) -> List[TemplateVariable]:
    """Prompt user to assign variable names to selected elements.

    Args:
        elements: List of selected elements.

    Returns:
        List of TemplateVariable objects.
    """
    variables = []

    for elem in elements:
        print(f'\nEnter variable name for "{elem.content}": ', end='')
        try:
            name = input().strip()
        except EOFError:
            name = f"var_{elem.index}"

        if not name:
            name = f"var_{elem.index}"

        # Sanitize variable name
        name = re.sub(r'[^a-zA-Z0-9_]', '_', name)

        variables.append(TemplateVariable(
            name=name,
            location=elem.location,
            position=elem.position,
            original_content=elem.content,
            syntax="both"
        ))

    return variables


def _run_has_drawing(run) -> bool:
    """Check if a run contains a drawing element (image).

    Args:
        run: The run to check.

    Returns:
        True if the run contains a drawing element.
    """
    # Check for w:drawing elements in the run's XML
    drawing_elements = run._r.findall(qn('w:drawing'))
    return len(drawing_elements) > 0


def _replace_paragraph_text_preserve_images(para, placeholder: str) -> None:
    """Replace text in a paragraph while preserving images/drawings.

    This function clears text from runs that don't contain drawings,
    and sets the placeholder in the first text-only run.

    Args:
        para: The paragraph to modify.
        placeholder: The placeholder text to insert.
    """
    first_text_run = None

    for run in para.runs:
        if _run_has_drawing(run):
            # Preserve runs with drawings - don't touch them
            continue
        else:
            # This is a text-only run
            if first_text_run is None:
                first_text_run = run
                run.text = placeholder
            else:
                # Clear subsequent text runs
                run.text = ''

    # If no text runs found, add a new one
    if first_text_run is None and para.runs:
        # Check if we can add text to an existing run
        para.add_run(placeholder)
    elif first_text_run is None:
        para.add_run(placeholder)


def replace_with_placeholder(
    doc: Document,
    element: DocumentElement,
    variable: TemplateVariable
) -> None:
    """Replace an element's content with a variable placeholder.

    Supports both <variable> and {{variable}} syntax by using {{variable}}.
    Preserves images and drawings in headers/footers.

    Args:
        doc: The document to modify.
        element: The element to replace.
        variable: The variable definition.
    """
    # Create placeholder using double brace syntax as the canonical form
    placeholder = f"{{{{{variable.name}}}}}"

    if element.location == 'cover' or element.location == 'body':
        # Replace in body paragraphs
        if element.paragraph_index < len(doc.paragraphs):
            para = doc.paragraphs[element.paragraph_index]
            _replace_paragraph_text_preserve_images(para, placeholder)

    elif element.location == 'header':
        for section in doc.sections:
            for para in section.header.paragraphs:
                if element.content.rstrip('...') in para.text:
                    _replace_paragraph_text_preserve_images(para, placeholder)
                    break

    elif element.location == 'header_first':
        for section in doc.sections:
            if section.different_first_page_header_footer:
                for para in section.first_page_header.paragraphs:
                    if element.content.rstrip('...') in para.text:
                        _replace_paragraph_text_preserve_images(para, placeholder)
                        break

    elif element.location == 'footer':
        for section in doc.sections:
            for para in section.footer.paragraphs:
                if element.content.rstrip('...') in para.text:
                    _replace_paragraph_text_preserve_images(para, placeholder)
                    break

    elif element.location == 'footer_first':
        for section in doc.sections:
            if section.different_first_page_header_footer:
                for para in section.first_page_footer.paragraphs:
                    if element.content.rstrip('...') in para.text:
                        _replace_paragraph_text_preserve_images(para, placeholder)
                        break


def remove_body_content(doc: Document, keep_first_n: int = 0) -> None:
    """Remove body content after the cover page.

    This removes all paragraphs and tables after the cover page boundary,
    not just clearing their content. The cover page paragraphs are preserved.

    Args:
        doc: The document to modify.
        keep_first_n: Number of paragraphs to keep (for cover page).
    """
    body = doc.element.body
    paragraphs = body.findall(qn('w:p'))
    tables = body.findall(qn('w:tbl'))

    # Remove paragraphs after the cover page
    for i, p in enumerate(paragraphs):
        if i >= keep_first_n:
            body.remove(p)

    # Remove all tables (they come after cover page typically)
    for tbl in tables:
        body.remove(tbl)


def strip_body_content(doc: Document, keep_first_n: int = 0) -> None:
    """Remove body content while preserving structure and styles.

    DEPRECATED: Use remove_body_content() instead for proper content removal.

    Args:
        doc: The document to modify.
        keep_first_n: Number of paragraphs to keep (for cover page).
    """
    # Remove paragraphs after the cover page section
    body = doc.element.body
    paragraphs = body.findall(qn('w:p'))

    for i, p in enumerate(paragraphs):
        if i >= keep_first_n:
            # Clear paragraph content but keep the element
            for child in list(p):
                if child.tag != qn('w:pPr'):  # Keep paragraph properties
                    p.remove(child)


def generate_manifest(
    template_name: str,
    source_file: str,
    variables: List[TemplateVariable],
    output_path: Path,
    cover_page_enabled: bool = False
) -> None:
    """Generate the template manifest.xml file.

    Args:
        template_name: Name of the template.
        source_file: Original source file name.
        variables: List of defined variables.
        output_path: Path to save the manifest.
        cover_page_enabled: Whether the template includes a cover page.
    """
    from datetime import timezone

    root = ET.Element('template')
    root.set('name', template_name)
    root.set('version', '1.0')

    source = ET.SubElement(root, 'source')
    source.text = source_file

    created = ET.SubElement(root, 'created')
    created.text = datetime.now(timezone.utc).isoformat().replace('+00:00', 'Z')

    features = ET.SubElement(root, 'features')

    # Cover page feature flag
    cover_elem = ET.SubElement(features, 'cover-page')
    cover_elem.set('enabled', 'true' if cover_page_enabled else 'false')

    # Variables feature flag
    if variables:
        vars_elem = ET.SubElement(features, 'variables')
        vars_elem.set('enabled', 'true')

    # Group variables by location
    locations = {}
    for var in variables:
        if var.location not in locations:
            locations[var.location] = []
        locations[var.location].append(var)

    for location, vars_list in locations.items():
        loc_elem = ET.SubElement(features, location.replace('_', '-'))
        loc_elem.set('enabled', 'true')

        for var in vars_list:
            var_elem = ET.SubElement(loc_elem, 'variable')
            var_elem.set('name', var.name)
            var_elem.set('location', f"{location}/{var.position}")
            var_elem.set('syntax', var.syntax)

    # Pretty print and save
    rough_string = ET.tostring(root, encoding='unicode')
    reparsed = minidom.parseString(rough_string)
    formatted = reparsed.toprettyxml(indent="  ")

    output_path.write_text(formatted, encoding='utf-8')


def generate_variables_xml(
    variables: List[TemplateVariable],
    output_path: Path
) -> None:
    """Generate the template variables.xml file.

    The original content of each variable becomes its default value,
    which is used when the variable is not provided or in non-interactive mode.

    Args:
        variables: List of defined variables.
        output_path: Path to save the variables file.
    """
    root = ET.Element('variables')
    root.set('version', '1.0')

    for var in variables:
        var_elem = ET.SubElement(root, 'variable')
        var_elem.set('name', var.name)

        # Default value is the original content (without truncation ellipsis)
        default = ET.SubElement(var_elem, 'default')
        original_text = var.original_content.rstrip('...')
        default.text = original_text

        # Description based on location and position
        description = ET.SubElement(var_elem, 'description')
        description.text = f"Variable in {var.location} ({var.position})"

        # Location info for reference
        location = ET.SubElement(var_elem, 'location')
        location.text = f"{var.location}/{var.position}"

    # Pretty print and save
    rough_string = ET.tostring(root, encoding='unicode')
    reparsed = minidom.parseString(rough_string)
    formatted = reparsed.toprettyxml(indent="  ")

    output_path.write_text(formatted, encoding='utf-8')


@dataclass
class VariableCandidate:
    """A potential variable candidate found during document analysis."""
    id: str  # e.g., "A1", "B1", "C2"
    category: str  # e.g., "textbox", "header", "header_first", "footer"
    category_label: str  # e.g., "Text Box", "Default Header"
    content: str  # The actual text content
    suggested_name: str  # Suggested variable name
    description: str  # Human-readable description


@dataclass
class DocumentAnalysis:
    """Result of analyzing a document for potential template variables."""
    source_file: str
    candidates: List[VariableCandidate]
    has_cover_page: bool
    has_different_first_page_header: bool
    paragraph_count: int
    cover_boundary: Optional[int]  # Detected cover page boundary

    def get_candidates_by_category(self, category: str) -> List[VariableCandidate]:
        """Get candidates filtered by category."""
        return [c for c in self.candidates if c.category == category]

    def print_summary(self) -> None:
        """Print a formatted summary of available variables."""
        print("=" * 60)
        print("AVAILABLE ELEMENTS FOR VARIABLES")
        print("=" * 60)

        categories = {}
        for c in self.candidates:
            if c.category not in categories:
                categories[c.category] = []
            categories[c.category].append(c)

        category_order = ['textbox', 'header', 'header_first', 'footer', 'footer_first']
        for cat in category_order:
            if cat in categories:
                items = categories[cat]
                label = items[0].category_label
                print(f"\n[{items[0].id[0]}] {label.upper()}")
                print("-" * 40)
                for item in items:
                    print(f"  {item.id}. \"{item.content}\"")

        print("\n" + "=" * 60)


def analyze_document_for_variables(source_docx: Path) -> DocumentAnalysis:
    """Analyze a DOCX document to find all potential template variables.

    This function examines a document and identifies elements that can be
    converted to template variables, including:
    - Text boxes (commonly used for cover page titles)
    - Headers (default and first page)
    - Footers (default and first page)
    - Cover page paragraphs

    Args:
        source_docx: Path to the DOCX file to analyze.

    Returns:
        DocumentAnalysis object containing all found variable candidates.

    Example:
        >>> analysis = analyze_document_for_variables(Path("proposal.docx"))
        >>> analysis.print_summary()
        >>> for candidate in analysis.candidates:
        ...     print(f"{candidate.id}: {candidate.content}")
    """
    source_docx = Path(source_docx)
    if not source_docx.exists():
        raise FileNotFoundError(f"Source file not found: {source_docx}")

    doc = Document(source_docx)
    candidates = []

    # Detect cover page boundary (first page break)
    cover_boundary = None
    for i, para in enumerate(doc.paragraphs[:30]):
        if has_page_break(para):
            cover_boundary = i
            break

    # Check for different first page header/footer
    has_different_first = any(
        section.different_first_page_header_footer
        for section in doc.sections
    )

    # === A. Text Box Elements ===
    textboxes = extract_textboxes_from_docx(source_docx)
    if textboxes:
        tb = textboxes[0]
        suggested_names = ['document_type', 'project_description', 'client_name', 'version']
        for i, line in enumerate(tb.paragraphs):
            suggested = suggested_names[i] if i < len(suggested_names) else f'textbox_line_{i+1}'
            candidates.append(VariableCandidate(
                id=f"A{i+1}",
                category='textbox',
                category_label='Text Box (Cover Page Title)',
                content=line.strip(),
                suggested_name=suggested,
                description=f"Text box line {i+1}"
            ))

    # === B. Default Header Elements ===
    header_elements = find_header_elements(doc)
    for elem in header_elements:
        candidates.append(VariableCandidate(
            id=f"B{elem.index}",
            category='header',
            category_label='Default Header (Normal Pages)',
            content=elem.content,
            suggested_name='project_name' if elem.index == 1 else f'header_{elem.index}',
            description=f"Header element {elem.index}"
        ))

    # === C. First Page Header Elements ===
    if has_different_first:
        first_header_elements = find_first_page_header_elements(doc)
        # Get full text from headers
        for section in doc.sections:
            if section.different_first_page_header_footer:
                paras = section.first_page_header.paragraphs
                for i, elem in enumerate(first_header_elements):
                    full_text = paras[i].text.strip() if i < len(paras) else elem.content
                    suggested = 'company_info' if i == 0 else 'date_contact' if i == 1 else f'header_first_{i+1}'
                    candidates.append(VariableCandidate(
                        id=f"C{elem.index}",
                        category='header_first',
                        category_label='First Page Header (Cover Page)',
                        content=full_text,
                        suggested_name=suggested,
                        description=f"First page header line {elem.index} (has logo)" if i == 0 else f"First page header line {elem.index}"
                    ))
                break

    # === D. Default Footer Elements ===
    footer_elements = find_footer_elements(doc)
    for elem in footer_elements:
        candidates.append(VariableCandidate(
            id=f"D{elem.index}",
            category='footer',
            category_label='Default Footer',
            content=elem.content,
            suggested_name=f'footer_{elem.index}',
            description=f"Footer element {elem.index} (usually page number)"
        ))

    # === E. First Page Footer Elements ===
    if has_different_first:
        first_footer_elements = find_first_page_footer_elements(doc)
        for elem in first_footer_elements:
            candidates.append(VariableCandidate(
                id=f"E{elem.index}",
                category='footer_first',
                category_label='First Page Footer',
                content=elem.content,
                suggested_name=f'footer_first_{elem.index}',
                description=f"First page footer element {elem.index}"
            ))

    return DocumentAnalysis(
        source_file=str(source_docx),
        candidates=candidates,
        has_cover_page=cover_boundary is not None,
        has_different_first_page_header=has_different_first,
        paragraph_count=len(doc.paragraphs),
        cover_boundary=cover_boundary
    )


def create_template(
    source_docx: Path,
    template_name: str,
    templates_dir: Optional[Path] = None,
    variables: Optional[Dict[str, str]] = None,
    interactive: bool = True,
    include_cover_page: bool = True,
    cover_boundary: Optional[int] = None
) -> Path:
    """Create a reusable template from a DOCX file.

    When creating a template:
    - Body content after the cover page is removed
    - Cover page content is preserved (up to the specified boundary)
    - Text boxes on cover page are scanned for variables
    - Variable placeholders replace original values (e.g., "John Doe" → <author>)
    - Original values are stored as defaults in variables.xml

    Args:
        source_docx: Path to the source DOCX file.
        template_name: Name for the template.
        templates_dir: Directory to store templates (default: ./templates).
        variables: Pre-defined variables for non-interactive mode.
        interactive: Whether to prompt for variable selection.
        include_cover_page: Whether to preserve cover page in template.
        cover_boundary: Paragraph index (0-based) where cover ends.
                       If None and interactive, user will be prompted.

    Returns:
        Path to the created template directory.
    """
    source_docx = Path(source_docx)
    if not source_docx.exists():
        raise FileNotFoundError(f"Source file not found: {source_docx}")

    # Set up template directory
    if templates_dir is None:
        templates_dir = Path('./templates')
    templates_dir = Path(templates_dir)

    template_path = templates_dir / template_name
    template_path.mkdir(parents=True, exist_ok=True)

    # Load the document
    doc = Document(source_docx)

    # Determine cover page boundary
    if include_cover_page:
        if cover_boundary is None and interactive:
            cover_boundary = prompt_cover_boundary(doc)
        elif cover_boundary is None:
            cover_boundary = 4  # Default: first 5 paragraphs
    else:
        cover_boundary = -1  # No cover page

    # Find potential variable elements
    cover_elements = find_cover_page_elements(doc, cover_boundary)
    header_elements = find_header_elements(doc)
    first_header_elements = find_first_page_header_elements(doc)
    footer_elements = find_footer_elements(doc)
    first_footer_elements = find_first_page_footer_elements(doc)

    # Extract text boxes (common for cover pages)
    textboxes = extract_textboxes_from_docx(source_docx)

    template_variables = []
    textbox_replacements = {}

    if interactive:
        # Interactive variable selection for paragraphs
        selected_cover = prompt_variable_selection(cover_elements, "cover page paragraphs")
        selected_header = prompt_variable_selection(header_elements, "default header")
        selected_first_header = prompt_variable_selection(first_header_elements, "first page (cover) header")
        selected_footer = prompt_variable_selection(footer_elements, "default footer")
        selected_first_footer = prompt_variable_selection(first_footer_elements, "first page (cover) footer")

        all_selected = selected_cover + selected_header + selected_first_header + selected_footer + selected_first_footer

        if all_selected:
            template_variables = prompt_variable_names(all_selected)

            # Replace selected elements with placeholders
            for elem, var in zip(all_selected, template_variables):
                replace_with_placeholder(doc, elem, var)

        # Interactive variable selection for text boxes
        if textboxes:
            textbox_selections, textbox_replacements = prompt_textbox_variable_selection(textboxes)
            if textbox_selections:
                textbox_vars = textbox_selections_to_variables(textbox_selections)
                template_variables.extend(textbox_vars)
    else:
        # Non-interactive mode: use provided variables
        if variables:
            for var_name, location in variables.items():
                template_variables.append(TemplateVariable(
                    name=var_name,
                    location=location.split('/')[0] if '/' in location else 'body',
                    position=location.split('/')[1] if '/' in location else location,
                    original_content='',
                    syntax='both'
                ))

    # Extract and save styles
    styles = extract_styles(source_docx)
    table_styles = extract_table_styles(source_docx)
    save_styles(styles, table_styles, template_path / 'styles.xml')

    # Remove body content after cover page
    if include_cover_page and cover_boundary >= 0:
        remove_body_content(doc, keep_first_n=cover_boundary + 1)
    else:
        # Remove all body content if no cover page
        remove_body_content(doc, keep_first_n=0)

    # Save template document (without text box replacements first)
    doc.save(template_path / 'template.docx')

    # Apply text box replacements if any
    if textbox_replacements:
        replace_text_in_textbox(
            template_path / 'template.docx',
            template_path / 'template.docx',
            textbox_replacements
        )

    # Generate manifest with cover page feature flag
    generate_manifest(
        template_name,
        source_docx.name,
        template_variables,
        template_path / 'manifest.xml',
        cover_page_enabled=include_cover_page and cover_boundary >= 0
    )

    # Generate variables.xml with defaults from original content
    generate_variables_xml(template_variables, template_path / 'variables.xml')

    # Extract fonts from source document
    fonts_extracted = _extract_template_fonts(source_docx, template_path, templates_dir)

    print(f"\nTemplate created: {template_path}/")
    if include_cover_page and cover_boundary >= 0:
        print(f"  Cover page: paragraphs 1-{cover_boundary + 1} preserved")
    if textboxes:
        print(f"  Text boxes found: {len(textboxes)}")
    print(f"  Variables: {len(template_variables)}")
    if fonts_extracted:
        print(f"  Fonts extracted: {len(fonts_extracted)}")

    return template_path


def _extract_template_fonts(
    source_docx: Path,
    template_path: Path,
    templates_dir: Path
) -> List[str]:
    """Extract fonts from source document to template fonts directory.

    If fonts are embedded in the source, extracts and deobfuscates them.
    If fonts are not embedded, searches default fonts and system fonts.

    Args:
        source_docx: Path to the source DOCX file.
        template_path: Path to the template directory.
        templates_dir: Path to the templates directory (for default fonts).

    Returns:
        List of font names that were extracted.
    """
    fonts_dir = template_path / 'fonts'
    extracted_fonts = []

    # Get fonts used in the document
    doc_fonts = extract_fonts_from_docx(source_docx)

    if not doc_fonts:
        return extracted_fonts

    # Check which fonts are embedded
    embedded_font_names = [name for name, info in doc_fonts.items() if info['embedded']]
    non_embedded_font_names = [name for name, info in doc_fonts.items() if not info['embedded']]

    # Extract embedded fonts
    if embedded_font_names:
        fonts_dir.mkdir(parents=True, exist_ok=True)
        extracted_paths = extract_embedded_fonts(source_docx, fonts_dir)
        if extracted_paths:
            for path in extracted_paths:
                font_name = get_font_name_from_file(path)
                if font_name:
                    extracted_fonts.append(font_name)

    # Search for non-embedded fonts in default fonts directory and system
    if non_embedded_font_names:
        default_fonts_dir = templates_dir / 'default' / 'fonts'
        search_dirs = []
        if default_fonts_dir.exists():
            search_dirs.append(default_fonts_dir)

        for font_name in non_embedded_font_names:
            # Skip common system fonts that don't need embedding
            if _is_common_system_font(font_name):
                continue

            # Find font files
            found = find_font_file(font_name, search_dirs, include_system=True)
            if found:
                fonts_dir.mkdir(parents=True, exist_ok=True)
                # Copy font files to template
                for variant, font_path in found.items():
                    dest_name = f"{font_name.replace(' ', '')}-{variant.capitalize()}{font_path.suffix}"
                    dest_path = fonts_dir / dest_name
                    if not dest_path.exists():
                        shutil.copy2(font_path, dest_path)
                extracted_fonts.append(font_name)
            else:
                print(f"  Warning: Could not find font '{font_name}' for embedding")

    return extracted_fonts


def _is_common_system_font(font_name: str) -> bool:
    """Check if a font is a common system font that doesn't need embedding.

    Args:
        font_name: Font family name.

    Returns:
        True if the font is commonly available on most systems.
    """
    common_fonts = {
        'arial', 'times new roman', 'times', 'courier new', 'courier',
        'verdana', 'georgia', 'trebuchet ms', 'impact', 'comic sans ms',
        'calibri', 'cambria', 'consolas', 'segoe ui',
        'helvetica', 'helvetica neue',
        'liberation sans', 'liberation serif', 'liberation mono',
        'dejavu sans', 'dejavu serif', 'dejavu sans mono',
    }
    return font_name.lower() in common_fonts
