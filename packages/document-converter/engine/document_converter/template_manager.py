"""Template management for DOCX documents."""
from pathlib import Path
from typing import Optional, List, Dict, Any
from xml.etree import ElementTree as ET
from docx import Document
import shutil
import re


def _find_project_root() -> Path:
    """Find the project root by looking for marker files.

    Searches up from the package location for directories containing
    .git, CLAUDE.md, or a templates/ directory.

    Returns:
        Path to the project root, or current working directory if not found.
    """
    # Start from the package directory
    current = Path(__file__).parent.resolve()

    # Walk up the directory tree
    for _ in range(10):  # Limit search depth
        # Check for project root markers
        if (current / '.git').exists() or (current / 'CLAUDE.md').exists():
            return current
        # Check if templates directory exists here
        if (current / 'templates').is_dir():
            return current
        # Move up one level
        parent = current.parent
        if parent == current:
            break
        current = parent

    # Fallback to current working directory
    return Path.cwd()


# Default templates directory (relative to project root)
DEFAULT_TEMPLATES_DIR = _find_project_root() / 'templates'


def resolve_templates_dir(templates_dir: Optional[Path] = None) -> Path:
    """Resolve the templates directory path.

    Args:
        templates_dir: Optional custom templates directory.

    Returns:
        Resolved Path to templates directory.
    """
    if templates_dir is not None:
        return Path(templates_dir)
    return DEFAULT_TEMPLATES_DIR


def list_templates(templates_dir: Optional[Path] = None) -> List[Dict[str, str]]:
    """List all available templates.

    Args:
        templates_dir: Optional custom templates directory.

    Returns:
        List of dictionaries with template info (name, description, path).
    """
    templates_path = resolve_templates_dir(templates_dir)

    if not templates_path.exists():
        return []

    templates = []
    for item in templates_path.iterdir():
        if item.is_dir():
            manifest_path = item / 'manifest.xml'
            template_info = {
                'name': item.name,
                'description': '',
                'path': str(item),
                'has_variables': False
            }

            # Try to read manifest for more info
            if manifest_path.exists():
                try:
                    tree = ET.parse(manifest_path)
                    root = tree.getroot()

                    # Get description if available
                    desc = root.find('description')
                    if desc is not None and desc.text:
                        template_info['description'] = desc.text

                    # Check for variables
                    features = root.find('features')
                    if features is not None:
                        for child in features:
                            if child.findall('variable'):
                                template_info['has_variables'] = True
                                break
                except Exception:
                    pass

            # Check if template.docx exists
            if (item / 'template.docx').exists():
                templates.append(template_info)

    return templates


def load_template(
    template_name: str,
    templates_dir: Optional[Path] = None
) -> Dict[str, Any]:
    """Load a template by name.

    Args:
        template_name: Name of the template to load.
        templates_dir: Optional custom templates directory.

    Returns:
        Dictionary containing template information and paths.

    Raises:
        FileNotFoundError: If template doesn't exist.
    """
    templates_path = resolve_templates_dir(templates_dir)
    template_path = templates_path / template_name

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_name}")

    template_docx = template_path / 'template.docx'
    if not template_docx.exists():
        raise FileNotFoundError(f"Template DOCX not found: {template_docx}")

    result = {
        'name': template_name,
        'path': template_path,
        'template_docx': template_docx,
        'fonts_dir': template_path / 'fonts' if (template_path / 'fonts').exists() else None,
        'variables': [],
        'styles': None,
        'has_cover_page': False,
        'has_textbox_variables': False
    }

    # Load manifest to check for cover page and features
    manifest_path = template_path / 'manifest.xml'
    if manifest_path.exists():
        try:
            tree = ET.parse(manifest_path)
            root = tree.getroot()
            features = root.find('features')
            if features is not None:
                # Check cover-page
                cover_page = features.find('cover-page')
                if cover_page is not None:
                    result['has_cover_page'] = cover_page.get('enabled', 'false').lower() == 'true'
                # Check textbox
                textbox = features.find('textbox')
                if textbox is not None:
                    result['has_textbox_variables'] = textbox.get('enabled', 'false').lower() == 'true'
        except Exception:
            pass

    # Load variables
    variables_path = template_path / 'variables.xml'
    if variables_path.exists():
        try:
            tree = ET.parse(variables_path)
            root = tree.getroot()
            for var in root.findall('variable'):
                result['variables'].append({
                    'name': var.get('name', ''),
                    'syntax': var.get('syntax', 'both'),
                    'location': var.findtext('location', ''),
                    'original': var.findtext('original-content', ''),
                    'default': var.findtext('default', '')
                })
        except Exception:
            pass

    # Load styles path
    styles_path = template_path / 'styles.xml'
    if styles_path.exists():
        result['styles'] = styles_path

    return result


def replace_variable(text: str, var_name: str, value: str) -> str:
    """Replace a variable in text, supporting both syntaxes.

    Replaces both <var_name> and {{var_name}} with the given value.

    Args:
        text: The text containing variables.
        var_name: The variable name to replace.
        value: The value to substitute.

    Returns:
        Text with variables replaced.
    """
    # Replace angle bracket syntax: <var_name>
    text = text.replace(f'<{var_name}>', value)

    # Replace double brace syntax: {{var_name}}
    text = text.replace(f'{{{{{var_name}}}}}', value)

    return text


def replace_all_variables(text: str, variables: Dict[str, str]) -> str:
    """Replace all variables in text.

    Args:
        text: The text containing variables.
        variables: Dictionary mapping variable names to values.

    Returns:
        Text with all variables replaced.
    """
    for var_name, value in variables.items():
        text = replace_variable(text, var_name, value)
    return text


def _run_has_special_elements(run) -> bool:
    """Check if a run contains special elements that would be lost if text is set.

    Setting run.text replaces the entire run content, which would lose
    page breaks, drawings, and other special elements.

    Args:
        run: The Run object to check.

    Returns:
        True if the run contains special elements that should be preserved.
    """
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Check for page breaks
    for br in run._element.iter(f'{w_ns}br'):
        br_type = br.get(f'{w_ns}type')
        if br_type == 'page':
            return True

    # Check for drawings (images, shapes)
    drawing_ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    if run._element.find(f'.//{drawing_ns}drawing') is not None:
        return True
    if run._element.find(f'.//w:drawing', {'w': w_ns[1:-1]}) is not None:
        return True

    return False


def _safe_replace_run_text(run, variables: Dict[str, str]) -> None:
    """Safely replace variables in a run's text without losing special elements.

    Only modifies the run if:
    1. The run has text
    2. The text actually changes after replacement
    3. The run doesn't contain special elements (page breaks, etc.)

    Args:
        run: The Run object to modify.
        variables: Dictionary mapping variable names to values.
    """
    if not run.text:
        return

    new_text = replace_all_variables(run.text, variables)

    # Only modify if text actually changes
    if new_text == run.text:
        return

    # Don't modify runs with special elements
    if _run_has_special_elements(run):
        return

    run.text = new_text


def replace_placeholders_in_document(
    doc: Document,
    variables: Dict[str, str]
) -> None:
    """Replace variable placeholders throughout the document.

    Supports both <variable> and {{variable}} syntax.

    Note: This function is careful not to modify runs that contain special
    elements like page breaks, as setting run.text would lose those elements.

    Args:
        doc: The Document to modify.
        variables: Dictionary mapping variable names to values.
    """
    # Replace in body paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            _safe_replace_run_text(run, variables)

    # Replace in headers
    for section in doc.sections:
        for para in section.header.paragraphs:
            for run in para.runs:
                _safe_replace_run_text(run, variables)

    # Replace in footers
    for section in doc.sections:
        for para in section.footer.paragraphs:
            for run in para.runs:
                _safe_replace_run_text(run, variables)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _safe_replace_run_text(run, variables)


def replace_placeholders_in_header(doc: Document, document_name: str, document_title: str) -> None:
    """Replace placeholder text in document headers.

    Supports both <placeholder> and {{placeholder}} syntax.

    Args:
        doc: Document object.
        document_name: Text to replace document name placeholder.
        document_title: Text to replace document title placeholder.
    """
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            # Check for both syntax forms
            has_placeholder = (
                '<document name>' in paragraph.text or
                '<document title>' in paragraph.text or
                '{{document name}}' in paragraph.text or
                '{{document_name}}' in paragraph.text or
                '{{document title}}' in paragraph.text or
                '{{document_title}}' in paragraph.text
            )
            if has_placeholder:
                # Replace placeholders in paragraph text (both syntaxes)
                text = paragraph.text
                # Angle bracket syntax
                text = text.replace('<document name>', document_name)
                text = text.replace('<document title>', document_title)
                # Double brace syntax (with space)
                text = text.replace('{{document name}}', document_name)
                text = text.replace('{{document title}}', document_title)
                # Double brace syntax (with underscore)
                text = text.replace('{{document_name}}', document_name)
                text = text.replace('{{document_title}}', document_title)

                # Clear existing runs and add new text
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = text
                else:
                    paragraph.add_run(text)


def apply_template(
    output_path: Path,
    template_path: Path,
    document_name: Optional[str] = None,
    document_title: Optional[str] = None
) -> None:
    """Apply template to generated DOCX file.

    This function:
    1. Opens the generated DOCX file
    2. Replaces placeholders in headers
    3. Preserves footers from template

    Args:
        output_path: Path to the generated DOCX file.
        template_path: Path to the template DOCX file.
        document_name: Text for document name placeholder.
        document_title: Text for document title placeholder.
    """
    if not template_path.exists():
        print(f"Warning: Template not found at {template_path}, skipping template application")
        return

    # Open the generated document
    doc = Document(output_path)
    template_doc = Document(template_path)

    # Get the first section from template (we'll apply this to all sections in generated doc)
    if len(template_doc.sections) == 0:
        print("Warning: Template has no sections, skipping template application")
        return

    template_section = template_doc.sections[0]

    # Copy header and footer from template to all sections in generated document
    for section in doc.sections:
        # Copy header
        section.header._element.clear()
        for element in template_section.header._element:
            section.header._element.append(element)

        # Copy footer
        section.footer._element.clear()
        for element in template_section.footer._element:
            section.footer._element.append(element)

    # Replace placeholders in headers
    if document_name or document_title:
        replace_placeholders_in_header(
            doc,
            document_name or '',
            document_title or ''
        )

    # Save the modified document
    doc.save(output_path)
