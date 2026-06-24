"""Text box extraction and manipulation for DOCX documents.

Handles text boxes (txbxContent elements) which are commonly used for
cover pages, callouts, and other formatted text blocks in Word documents.
"""
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import zipfile
import re


# XML namespaces used in DOCX
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'v': 'urn:schemas-microsoft-com:vml',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
}


@dataclass
class TextBoxElement:
    """Represents a text element within a text box."""
    index: int
    full_text: str
    paragraphs: List[str]
    location: str  # 'textbox'
    xpath: str  # XPath to the txbxContent element
    paragraph_index: int  # Index in document paragraphs (approximate)


def extract_textboxes_from_docx(doc_path: Path) -> List[TextBoxElement]:
    """Extract all text boxes from a DOCX file.

    Args:
        doc_path: Path to the DOCX file.

    Returns:
        List of TextBoxElement objects.
    """
    textboxes = []

    with zipfile.ZipFile(doc_path, 'r') as zf:
        with zf.open('word/document.xml') as f:
            tree = etree.parse(f)
            root = tree.getroot()

            # Find all txbxContent elements (text box content)
            # Two common patterns:
            # 1. Drawing-based: w:drawing/...wps:txbx/w:txbxContent
            # 2. VML-based: w:pict/v:rect/v:textbox/w:txbxContent

            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

            index = 1
            seen_texts = set()  # Avoid duplicates from MC alternateContent

            for txbx_content in root.iter(f'{w_ns}txbxContent'):
                paragraphs = []

                # Extract text from each paragraph in the text box
                for p_elem in txbx_content.findall(f'{w_ns}p'):
                    para_texts = []
                    for t_elem in p_elem.iter(f'{w_ns}t'):
                        if t_elem.text:
                            para_texts.append(t_elem.text)
                    if para_texts:
                        paragraphs.append(''.join(para_texts))

                if paragraphs:
                    full_text = '\n'.join(paragraphs)

                    # Skip duplicates (MC alternateContent creates copies)
                    if full_text in seen_texts:
                        continue
                    seen_texts.add(full_text)

                    # Try to find approximate paragraph index
                    para_idx = _find_paragraph_index(root, txbx_content, w_ns)

                    textboxes.append(TextBoxElement(
                        index=index,
                        full_text=full_text,
                        paragraphs=paragraphs,
                        location='textbox',
                        xpath=_get_xpath(txbx_content),
                        paragraph_index=para_idx
                    ))
                    index += 1

    return textboxes


def _find_paragraph_index(root, txbx_content, w_ns: str) -> int:
    """Find the approximate paragraph index containing the text box."""
    # Walk up to find the containing paragraph
    parent = txbx_content.getparent()
    while parent is not None:
        if parent.tag == f'{w_ns}p':
            # Count paragraphs before this one
            body = root.find(f'{w_ns}body')
            if body is not None:
                for i, p in enumerate(body.findall(f'{w_ns}p')):
                    if p is parent:
                        return i
            break
        parent = parent.getparent()
    return 0


def _get_xpath(element) -> str:
    """Get a simplified XPath for an element."""
    parts = []
    current = element
    while current is not None:
        tag = current.tag.split('}')[-1] if '}' in current.tag else current.tag
        parts.insert(0, tag)
        current = current.getparent()
    return '/' + '/'.join(parts[-5:])  # Last 5 elements


def display_textboxes(textboxes: List[TextBoxElement]) -> None:
    """Display text boxes for user selection.

    Args:
        textboxes: List of text box elements.
    """
    if not textboxes:
        print("\nNo text boxes found in document.")
        return

    print("\n=== Text Boxes ===")
    for tb in textboxes:
        print(f"\n[{tb.index}] Text Box (at paragraph ~{tb.paragraph_index + 1}):")
        for i, para in enumerate(tb.paragraphs):
            display_text = para[:60] + ('...' if len(para) > 60 else '')
            print(f"     Line {i+1}: \"{display_text}\"")


def escape_xml(text: str) -> str:
    """Escape special XML characters in text.

    Args:
        text: Text to escape.

    Returns:
        XML-escaped text.
    """
    return (text
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&apos;'))


def replace_fragmented_text(content: str, old_text: str, new_text: str) -> str:
    """Replace text that may be fragmented across multiple w:t elements.

    Word often splits text across multiple <w:t> elements due to formatting,
    spelling/grammar checks, and other processing. This function handles:
    - Simple contiguous text
    - Text split by XML tags like </w:t><w:proofErr.../><w:r><w:t>
    - Variable patterns like {{var_name}} split across elements

    Args:
        content: XML content.
        old_text: Text to find (may be split across elements).
        new_text: Replacement text.

    Returns:
        Modified XML content.
    """
    import re

    # First try simple replacement
    if old_text in content:
        return content.replace(old_text, new_text)

    # For variable patterns like {{var_name}}, build a more flexible regex
    # that allows ANY XML content between the text characters
    chars = list(old_text)
    pattern_parts = []
    for char in chars:
        # Escape regex special characters
        escaped_char = re.escape(char)
        pattern_parts.append(escaped_char)

    # Allow any XML tags/content between characters (including proofErr, formatting, etc.)
    # This pattern allows: </w:t>, any number of tags, then <w:t...>
    # The key is to be very permissive about what can appear between text characters
    xml_between = r'(?:</w:t>(?:[^<]*<[^>]*>)*[^<]*<w:t[^>]*>)?'
    pattern = xml_between.join(pattern_parts)

    # Find all matches
    matches = list(re.finditer(pattern, content, re.DOTALL))

    if not matches:
        # Try an even more aggressive approach for {{variable}} patterns
        # Match the opening {{ and closing }} with the variable name in between
        if old_text.startswith('{{') and old_text.endswith('}}'):
            var_name = old_text[2:-2].strip()
            # Build pattern: {{ ... var_name ... }}
            # Allow XML between {{ and var_name, and between var_name and }}
            aggressive_pattern = (
                r'\{\{\s*'  # Opening {{ with optional space
                r'(?:</w:t>(?:[^<]*<[^>]*>)*[^<]*<w:t[^>]*>)?'  # Allow XML breaks
                r'\s*' + re.escape(var_name.replace('_', '')) +  # var name without underscores
                r'|' + re.escape(var_name) +  # or exact var name
                r'(?:</w:t>(?:[^<]*<[^>]*>)*[^<]*<w:t[^>]*>)?'  # Allow XML breaks
                r'\s*\}\}'  # Closing }} with optional space
            )
            # This is too complex, let's try a different approach
            pass
        return content

    # Replace from end to start to maintain positions
    result = content
    for match in reversed(matches):
        result = result[:match.start()] + new_text + result[match.end():]

    return result


def replace_fragmented_variable(content: str, var_name: str, new_value: str) -> str:
    """Replace a {{variable}} pattern that may be fragmented across XML elements.

    This is specialized for handling Word's tendency to split variable placeholders
    like {{document_type}} into multiple <w:t> elements with proofErr tags.

    Strategy: Find runs of <w:t> elements that together contain the variable pattern,
    then replace the text content while preserving XML structure.

    Args:
        content: XML content.
        var_name: Variable name (without {{ }}).
        new_value: Replacement value.

    Returns:
        Modified XML content.
    """
    import re

    # First try simple replacement (variable is in a single <w:t> element)
    simple_pattern = '{{' + var_name + '}}'
    if simple_pattern in content:
        return content.replace(simple_pattern, new_value)

    # Also try with spaces around the variable name (Word sometimes adds spaces)
    spaced_pattern = '{{ ' + var_name + ' }}'
    if spaced_pattern in content:
        return content.replace(spaced_pattern, new_value)

    # For fragmented patterns, we need to find <w:t> elements that together
    # contain {{ var_name }}. The approach:
    # 1. Find all <w:t>...</w:t> elements
    # 2. Look for sequences where the combined text matches our pattern
    # 3. Replace the text in the first element, clear the others

    # Pattern to find <w:t> elements with their content
    wt_pattern = r'(<w:t[^>]*>)([^<]*)(</w:t>)'

    # Find all matches with positions
    matches = list(re.finditer(wt_pattern, content))

    if not matches:
        return content

    # Build a map of text content and positions
    # Look for sequences that form {{ var_name }}
    target = '{{ ' + var_name + ' }}'  # With spaces
    target_no_space = '{{' + var_name + '}}'  # Without spaces

    result = content
    replaced = True
    max_iterations = 100  # Safety limit
    iteration = 0

    while replaced and iteration < max_iterations:
        replaced = False
        iteration += 1
        matches = list(re.finditer(wt_pattern, result))

        for i in range(len(matches)):
            # Try to build the target string from consecutive matches
            combined_text = ''
            match_indices = []

            for j in range(i, min(i + 20, len(matches))):  # Look ahead up to 20 elements
                combined_text += matches[j].group(2)  # Add text content
                match_indices.append(j)

                # Check if we found our target
                found_target = None
                if target in combined_text:
                    found_target = target
                elif target_no_space in combined_text:
                    found_target = target_no_space

                if found_target:
                    # Found the pattern! Now replace carefully
                    pattern_start = combined_text.find(found_target)
                    pattern_end = pattern_start + len(found_target)

                    # Calculate which <w:t> elements contain parts of the pattern
                    char_pos = 0
                    elements_to_modify = []
                    for idx in match_indices:
                        elem_text = matches[idx].group(2)
                        elem_start = char_pos
                        elem_end = char_pos + len(elem_text)

                        # Check if this element overlaps with the pattern
                        if elem_end > pattern_start and elem_start < pattern_end:
                            elements_to_modify.append((idx, elem_start, elem_end))

                        char_pos = elem_end

                    if elements_to_modify:
                        # Build replacement: put new value in first element, empty others
                        # But we need to preserve any text before/after the pattern
                        first_idx = elements_to_modify[0][0]
                        last_idx = elements_to_modify[-1][0]

                        first_match = matches[first_idx]
                        last_match = matches[last_idx]

                        # Text before pattern in first element
                        first_elem_text = first_match.group(2)
                        first_elem_start_in_combined = sum(len(matches[k].group(2)) for k in match_indices if k < first_idx)
                        before_text = combined_text[first_elem_start_in_combined:pattern_start] if pattern_start > first_elem_start_in_combined else ''

                        # Text after pattern in last element
                        after_text = combined_text[pattern_end:]

                        # Build the new content for the range
                        # Keep XML structure intact - only modify text within <w:t> elements
                        new_content_parts = []
                        for idx in range(first_idx, last_idx + 1):
                            m = matches[idx]
                            if idx == first_idx:
                                # First element: before_text + new_value (+ after if same element)
                                if idx == last_idx:
                                    new_content_parts.append(m.group(1) + before_text + new_value + after_text + m.group(3))
                                else:
                                    new_content_parts.append(m.group(1) + before_text + new_value + m.group(3))
                            elif idx == last_idx:
                                # Last element: just after_text
                                new_content_parts.append(m.group(1) + after_text + m.group(3))
                            else:
                                # Middle elements: empty
                                new_content_parts.append(m.group(1) + m.group(3))

                        # Replace the range
                        start_pos = first_match.start()
                        end_pos = last_match.end()

                        # Get content between elements (XML tags, etc.)
                        inter_element_content = []
                        for idx in range(first_idx, last_idx):
                            inter_element_content.append(result[matches[idx].end():matches[idx+1].start()])

                        # Rebuild with new text but same structure
                        replacement = ''
                        for k, part in enumerate(new_content_parts):
                            replacement += part
                            if k < len(inter_element_content):
                                replacement += inter_element_content[k]

                        result = result[:start_pos] + replacement + result[end_pos:]
                        replaced = True
                        break

                    break  # Found pattern, move to next search

            if replaced:
                break  # Restart search from beginning

    return result


def replace_text_in_textbox(
    doc_path: Path,
    output_path: Path,
    replacements: Dict[str, str]
) -> None:
    """Replace text in text boxes within a DOCX file.

    Handles XML escaping for placeholder syntax like <variable>.
    Also handles text that is fragmented across multiple <w:t> elements.

    Args:
        doc_path: Path to the source DOCX file.
        output_path: Path to save the modified DOCX.
        replacements: Dictionary mapping old text to new text (e.g., placeholders).
    """
    import tempfile
    import os

    # Create a temporary directory
    with tempfile.TemporaryDirectory() as tmpdir:
        # Extract the DOCX
        with zipfile.ZipFile(doc_path, 'r') as zf:
            zf.extractall(tmpdir)

        # Modify document.xml and all header/footer files
        word_dir = Path(tmpdir) / 'word'
        xml_files = [word_dir / 'document.xml']

        # Dynamically find all header and footer files
        if word_dir.exists():
            for xml_file in word_dir.glob('header*.xml'):
                xml_files.append(xml_file)
            for xml_file in word_dir.glob('footer*.xml'):
                xml_files.append(xml_file)

        for xml_path in xml_files:
            if not xml_path.exists():
                continue

            with open(xml_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Perform replacements with XML escaping
            for old_text, new_text in replacements.items():
                # Escape both old and new text for XML matching
                # e.g., <var> in search becomes &lt;var&gt; to match the XML content
                escaped_new = escape_xml(new_text)

                # Check if this is a {{variable}} pattern
                if old_text.startswith('{{') and old_text.endswith('}}'):
                    var_name = old_text[2:-2]  # Extract variable name
                    # Use specialized variable replacement for fragmented patterns
                    content = replace_fragmented_variable(content, var_name, escaped_new)
                else:
                    # Use general fragmented text replacement
                    escaped_old = escape_xml(old_text)
                    content = replace_fragmented_text(content, escaped_old, escaped_new)

            with open(xml_path, 'w', encoding='utf-8') as f:
                f.write(content)

        # Repack the DOCX
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for root_dir, dirs, files in os.walk(tmpdir):
                for file in files:
                    file_path = Path(root_dir) / file
                    arcname = file_path.relative_to(tmpdir)
                    zf.write(file_path, arcname)


def extract_textbox_variables(
    doc_path: Path,
    interactive: bool = True
) -> Tuple[List[TextBoxElement], Dict[str, str]]:
    """Extract text boxes and optionally prompt for variable selection.

    Args:
        doc_path: Path to the DOCX file.
        interactive: Whether to prompt for user input.

    Returns:
        Tuple of (selected textboxes, replacements dict).
    """
    textboxes = extract_textboxes_from_docx(doc_path)

    if not textboxes:
        return [], {}

    if not interactive:
        return textboxes, {}

    display_textboxes(textboxes)

    print("\nWhich text box elements should be variables?")
    print("Enter as: textbox_index.line_index (e.g., '1.1,1.2,1.3' for lines 1-3 of textbox 1)")
    print("Or 'none' to skip: ", end='')

    try:
        response = input().strip().lower()
    except EOFError:
        return [], {}

    if response == 'none' or response == '':
        return [], {}

    selected = []
    replacements = {}

    # Parse selection like "1.1,1.2,1.3"
    for item in response.split(','):
        item = item.strip()
        if '.' in item:
            try:
                tb_idx, line_idx = item.split('.')
                tb_idx = int(tb_idx)
                line_idx = int(line_idx)

                # Find the textbox
                for tb in textboxes:
                    if tb.index == tb_idx and line_idx <= len(tb.paragraphs):
                        para_text = tb.paragraphs[line_idx - 1]

                        print(f"\nEnter variable name for \"{para_text[:50]}{'...' if len(para_text) > 50 else ''}\": ", end='')
                        try:
                            var_name = input().strip()
                        except EOFError:
                            var_name = f"var_{tb_idx}_{line_idx}"

                        if not var_name:
                            var_name = f"var_{tb_idx}_{line_idx}"

                        # Sanitize variable name
                        var_name = re.sub(r'[^a-zA-Z0-9_]', '_', var_name)

                        replacements[para_text] = f"{{{{{var_name}}}}}"
                        selected.append((tb, line_idx, var_name, para_text))
                        break
            except ValueError:
                continue

    return textboxes, replacements


def find_textbox_text(doc_path: Path, search_text: str) -> Optional[TextBoxElement]:
    """Find a text box containing specific text.

    Args:
        doc_path: Path to the DOCX file.
        search_text: Text to search for.

    Returns:
        TextBoxElement if found, None otherwise.
    """
    textboxes = extract_textboxes_from_docx(doc_path)

    for tb in textboxes:
        if search_text in tb.full_text:
            return tb
        for para in tb.paragraphs:
            if search_text in para:
                return tb

    return None
