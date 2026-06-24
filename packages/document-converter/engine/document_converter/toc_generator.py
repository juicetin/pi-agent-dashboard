"""Table of Contents generator for DOCX documents."""
from pathlib import Path
from typing import List, Tuple, Optional
from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips
import re

from .toc_style import TocStyleConfig, TocLevelStyle, TocConfig, TocMode, get_tab_leader_char


def create_toc_field(style_config: Optional[TocStyleConfig] = None):
    """Create a TOC field element that Word will recognize and update.

    Args:
        style_config: Style configuration for the TOC field.

    Returns:
        A paragraph element containing the TOC field.
    """
    if style_config is None:
        style_config = TocStyleConfig()

    # Create paragraph for TOC
    paragraph = OxmlElement('w:p')

    # Create run
    run = OxmlElement('w:r')

    # Create field character (begin)
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    # Build TOC field code based on style configuration
    # \o "1-N" = include heading levels 1-N
    # \h = hyperlinks
    # \z = hide tab leader and page numbers in Web Layout view
    # \u = use applied paragraph outline level
    # \n = no page numbers (if show_page_numbers is False)
    field_code_parts = [f'TOC \\o "1-{style_config.max_level}"']

    if style_config.hyperlinks:
        field_code_parts.append('\\h')

    field_code_parts.append('\\z \\u')

    if not style_config.show_page_numbers:
        field_code_parts.append('\\n')

    field_code = ' ' + ' '.join(field_code_parts) + ' '

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_code

    # Create field character (separate)
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    # Create placeholder text
    placeholder_run = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = 'Right-click and select "Update Field" to generate table of contents'
    placeholder_run.append(placeholder_text)

    # Create field character (end)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    # Assemble the field
    run.append(fldChar_begin)
    paragraph.append(run)

    run2 = OxmlElement('w:r')
    run2.append(instrText)
    paragraph.append(run2)

    run3 = OxmlElement('w:r')
    run3.append(fldChar_separate)
    paragraph.append(run3)

    paragraph.append(placeholder_run)

    run4 = OxmlElement('w:r')
    run4.append(fldChar_end)
    paragraph.append(run4)

    return paragraph


def create_toc_heading(text: str = "Tartalomjegyzék"):
    """Create a heading paragraph for the TOC.

    Args:
        text: The heading text (default: "Tartalomjegyzék" - Hungarian for Table of Contents)

    Returns:
        A paragraph element with heading styling.
    """
    paragraph = OxmlElement('w:p')

    # Paragraph properties
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.append(pStyle)

    # Add spacing after
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:after'), '240')  # 12pt after
    pPr.append(spacing)

    paragraph.append(pPr)

    # Run with text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    bold = OxmlElement('w:b')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '32')  # 16pt
    rPr.append(bold)
    rPr.append(sz)
    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    paragraph.append(run)

    return paragraph


def create_page_break():
    """Create a page break paragraph.

    Returns:
        A paragraph element containing a page break.
    """
    paragraph = OxmlElement('w:p')
    run = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run.append(br)
    paragraph.append(run)
    return paragraph


def find_first_page_break(body) -> tuple:
    """Find the first page break in the document body.

    Args:
        body: The document body element.

    Returns:
        Tuple of (index, element) where page break was found, or (0, None) if not found.
    """
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    for i, elem in enumerate(body):
        # Check for page break in paragraph
        for br in elem.iter(f'{w_ns}br'):
            if br.get(qn('w:type')) == 'page':
                return i, elem

        # Also check for section break (continuous or nextPage)
        for sectPr in elem.iter(f'{w_ns}sectPr'):
            sect_type = sectPr.find(f'{w_ns}type')
            if sect_type is not None and sect_type.get(qn('w:val')) in ('nextPage', 'oddPage', 'evenPage'):
                return i, elem

    return 0, None


def find_first_heading(doc: Document) -> Optional[int]:
    """Find the index of the first heading paragraph in the document.

    Args:
        doc: The Document object.

    Returns:
        Index of the first heading paragraph, or None if not found.
    """
    body = doc.element.body
    body_elements = list(body)

    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ''
        if style_name.startswith('Heading'):
            # Find the corresponding element index in body
            try:
                return body_elements.index(para._p)
            except ValueError:
                continue

    return None


def add_toc(
    docx_path: Path,
    toc_heading: str = "Tartalomjegyzék",
    add_page_break: bool = True,
    after_cover_page: bool = False,
    style_config: Optional[TocStyleConfig] = None
) -> None:
    """Add a Table of Contents to a DOCX document.

    The TOC is a Word field that will be updated when the document is opened
    in Word and the user clicks "Update Table" or presses F9.

    Args:
        docx_path: Path to the DOCX file.
        toc_heading: Heading text for the TOC (default: "Tartalomjegyzék").
        add_page_break: Whether to add a page break after the TOC (default: True).
        after_cover_page: If True, insert TOC after the first page break (cover page).
        style_config: Style configuration for TOC entries.
    """
    if style_config is None:
        style_config = TocStyleConfig()

    doc = Document(docx_path)

    # Get the document body element
    body = doc.element.body

    # Determine insertion point
    if after_cover_page:
        # Find first page break and insert after it
        page_break_idx, page_break_elem = find_first_page_break(body)
        if page_break_elem is not None:
            # Insert after the element containing the page break
            insert_after = page_break_elem
        else:
            # No page break found, insert at beginning
            insert_after = None
    else:
        insert_after = None

    # Create TOC elements
    toc_heading_elem = create_toc_heading(toc_heading)
    toc_field = create_toc_field(style_config)

    if insert_after is not None:
        # Insert after the cover page (after the page break element)
        next_elem = insert_after.getnext()

        # Insert TOC heading after the page break
        insert_after.addnext(toc_heading_elem)
        # Insert TOC field after heading
        toc_heading_elem.addnext(toc_field)

        # Add page break after TOC if requested
        if add_page_break:
            page_break = create_page_break()
            toc_field.addnext(page_break)
    else:
        # Insert at beginning (original behavior)
        first_element = body[0] if len(body) > 0 else None

        # Insert in reverse order (so they end up in correct order)
        if add_page_break:
            page_break = create_page_break()
            if first_element is not None:
                first_element.addprevious(page_break)
            else:
                body.append(page_break)
            first_element = page_break

        if first_element is not None:
            first_element.addprevious(toc_field)
        else:
            body.append(toc_field)

        # Insert heading before TOC field
        toc_field.addprevious(toc_heading_elem)

    doc.save(docx_path)


def extract_headings(doc: Document, max_level: int = 3) -> List[Tuple[int, str, str]]:
    """Extract headings from a DOCX document.

    Args:
        doc: The Document object.
        max_level: Maximum heading level to include (1-9).

    Returns:
        List of tuples: (level, text, bookmark_name)
    """
    headings = []
    bookmark_counter = 0
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''

        # Check if it's a heading style
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading ', '').replace('Heading', ''))
                if level <= max_level:
                    text = para.text.strip()
                    if text:
                        # Generate bookmark name
                        bookmark_name = f"_Toc{bookmark_counter:08d}"
                        bookmark_counter += 1
                        headings.append((level, text, bookmark_name))

                        # Add bookmark to the paragraph
                        _add_bookmark_to_paragraph(para, bookmark_name)
            except ValueError:
                pass

    return headings


def _add_bookmark_to_paragraph(para, bookmark_name: str) -> None:
    """Add a bookmark to a paragraph.

    Args:
        para: The paragraph to add the bookmark to.
        bookmark_name: The bookmark name.
    """
    # Create bookmark start
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(hash(bookmark_name) % 10000))
    bookmark_start.set(qn('w:name'), bookmark_name)

    # Create bookmark end
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(hash(bookmark_name) % 10000))

    # Insert at the beginning and end of the paragraph
    para._p.insert(0, bookmark_start)
    para._p.append(bookmark_end)


def _get_tab_leader_xml(leader_type: str) -> str:
    """Convert tab leader type to XML attribute value.

    Args:
        leader_type: Type of tab leader (dots, dashes, underline, none).

    Returns:
        XML attribute value for w:leader.
    """
    leaders = {
        'dots': 'dot',
        'dashes': 'hyphen',
        'underline': 'underscore',
        'none': 'none',
    }
    return leaders.get(leader_type.lower(), 'dot')


def create_static_toc_entry(
    level: int,
    text: str,
    bookmark_name: str,
    style_config: Optional[TocStyleConfig] = None,
    level_style: Optional[TocLevelStyle] = None
) -> OxmlElement:
    """Create a static TOC entry paragraph with hyperlink.

    Args:
        level: Heading level (1-9).
        text: The heading text.
        bookmark_name: The bookmark name to link to.
        style_config: Overall TOC style configuration.
        level_style: Specific style for this level.

    Returns:
        A paragraph element for the TOC entry.
    """
    if style_config is None:
        style_config = TocStyleConfig()

    paragraph = OxmlElement('w:p')

    # Paragraph properties
    pPr = OxmlElement('w:pPr')

    # Use TOC style
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), f'TOC{level}')
    pPr.append(pStyle)

    # Indentation - use level_style if provided, otherwise calculate from config
    indent = 0
    if level_style and level_style.indent is not None:
        indent = level_style.indent * 20  # points to twips
    else:
        indent = (level - 1) * style_config.indent_per_level * 20  # points to twips

    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(indent))
    pPr.append(ind)

    # Spacing
    if level_style:
        spacing = OxmlElement('w:spacing')
        if level_style.spacing_before is not None:
            spacing.set(qn('w:before'), str(level_style.spacing_before * 20))
        if level_style.spacing_after is not None:
            spacing.set(qn('w:after'), str(level_style.spacing_after * 20))
        if level_style.spacing_before is not None or level_style.spacing_after is not None:
            pPr.append(spacing)

    # Tab stops for dot leader (right-aligned at ~6 inches = 8640 twips)
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), _get_tab_leader_xml(style_config.tab_leader))
    tab.set(qn('w:pos'), '9072')  # ~6.3 inches
    tabs.append(tab)
    pPr.append(tabs)

    paragraph.append(pPr)

    # Create hyperlink to bookmark (if hyperlinks enabled)
    if style_config.hyperlinks:
        container = OxmlElement('w:hyperlink')
        container.set(qn('w:anchor'), bookmark_name)
        container.set(qn('w:history'), '1')
    else:
        container = paragraph  # Add run directly to paragraph

    # Run with text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Apply level-specific styling
    if level_style:
        if level_style.bold:
            bold = OxmlElement('w:b')
            rPr.append(bold)
        if level_style.italic:
            italic = OxmlElement('w:i')
            rPr.append(italic)
        if level_style.font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(level_style.font_size * 2))  # half-points
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(level_style.font_size * 2))
            rPr.append(szCs)
        if level_style.font_name:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), level_style.font_name)
            rFonts.set(qn('w:hAnsi'), level_style.font_name)
            rPr.append(rFonts)
        if level_style.color:
            color = OxmlElement('w:color')
            color.set(qn('w:val'), level_style.color)
            rPr.append(color)
    else:
        # Default: bold for level 1
        if level == 1:
            bold = OxmlElement('w:b')
            rPr.append(bold)

    run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    if style_config.hyperlinks:
        container.append(run)
        paragraph.append(container)
    else:
        paragraph.append(run)

    return paragraph


def add_static_toc(
    docx_path: Path,
    toc_heading: str = "Tartalomjegyzék",
    max_level: int = 3,
    add_page_break: bool = True,
    after_cover_page: bool = False,
    style_config: Optional[TocStyleConfig] = None
) -> None:
    """Add a static Table of Contents to a DOCX document.

    This creates actual TOC entries with hyperlinks to bookmarks,
    which works correctly when converting to PDF without requiring
    Word field updates.

    Args:
        docx_path: Path to the DOCX file.
        toc_heading: Heading text for the TOC.
        max_level: Maximum heading level to include (default: 3).
        add_page_break: Whether to add a page break after the TOC.
        after_cover_page: If True, insert TOC before the first heading
            (after cover page content like logos).
        style_config: Style configuration for TOC entries.
    """
    if style_config is None:
        style_config = TocStyleConfig()

    # Use max_level from style_config if not explicitly overridden
    effective_max_level = style_config.max_level if max_level == 3 else max_level

    doc = Document(docx_path)

    # Extract headings and add bookmarks
    headings = extract_headings(doc, effective_max_level)

    if not headings:
        # No headings found, save and return
        doc.save(docx_path)
        return

    # Get the document body element
    body = doc.element.body

    # Determine insertion point
    # Best approach: insert right before the first heading
    # This preserves any content between cover page and first heading (e.g., logo)
    first_heading_idx = find_first_heading(doc)

    # Build TOC elements in correct order:
    # 1. TOC heading
    # 2. TOC entries (in document order)
    # 3. Page break
    toc_elements = []

    # TOC heading first
    toc_elements.append(create_toc_heading(toc_heading))

    # TOC entries in document order
    for level, text, bookmark_name in headings:
        # Get level-specific style if defined
        level_style = style_config.level_styles.get(level)
        toc_entry = create_static_toc_entry(
            level, text, bookmark_name,
            style_config=style_config,
            level_style=level_style
        )
        toc_elements.append(toc_entry)

    # Page break after TOC
    if add_page_break:
        toc_elements.append(create_page_break())

    # Insert all elements at the correct position
    if first_heading_idx is not None:
        # Insert before the first heading
        for i, elem in enumerate(toc_elements):
            body.insert(first_heading_idx + i, elem)
    else:
        # No heading found, insert at beginning
        for i, elem in enumerate(toc_elements):
            body.insert(i, elem)

    doc.save(docx_path)
